# =============================================================
# StationDeck - Export Engine
# src/exporter.py
# =============================================================
# Generates PDF, DOCX, and XLSX report files.
#
# Usage:
#   engine = ExportEngine(station_config=station_config)
#   pdf_path  = engine.generate_pdf(metrics, sections, period_label)
#   docx_path = engine.generate_docx(metrics, sections, period_label)
#   xlsx_path = engine.generate_xlsx(metrics, daily_df, period_label)
#
# station_config is optional. When provided, output dirs and station
# identity come from it. When absent, falls back to config.settings.
# =============================================================

from pathlib import Path
from datetime import datetime
import re

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor, white, black
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from config.settings import (
    REPORTS_PDF_DIR,
    REPORTS_DOCX_DIR,
    STATION_NAME,
    STATION_LOCATION,
    REPORT_CURRENCY,
    REPORT_AUTHOR,
)

# ── Colour constants ──────────────────────────────────────────────────────────

BRAND_RED     = HexColor("#C8102E")
BRAND_DARK    = HexColor("#1A1A2E")
BRAND_LIGHT   = HexColor("#F5F5F5")
BRAND_MID     = HexColor("#E0E0E0")
TEXT_DARK     = HexColor("#222222")
SURPLUS_GREEN = HexColor("#1E7E34")
DEFICIT_RED   = HexColor("#C8102E")

W_RED  = RGBColor(0xC8, 0x10, 0x2E)
W_DARK = RGBColor(0x1A, 0x1A, 0x2E)
W_GREY = RGBColor(0xF5, 0xF5, 0xF5)
W_MID  = RGBColor(0xE0, 0xE0, 0xE0)

XL_DARK    = "1A1A2E"
XL_RED     = "C8102E"
XL_LIGHT   = "F5F5F5"
XL_WHITE   = "FFFFFF"
XL_MID     = "E0E0E0"
XL_GREEN   = "1E7E34"
XL_SUBHEAD = "2E4057"


# ── Period label sanitiser ────────────────────────────────────────────────────

def _clean_period(period_label):
    """
    Extract only 'Month YYYY' from the period label.

    Defends against upstream bugs where a date stamp gets concatenated
    onto the label, producing e.g. 'April 2026 20260520' instead of
    'April 2026'.

    Strategy: keep only the first two whitespace-separated tokens
    (the month name and the 4-digit year). Everything after is discarded.

    Examples:
        'April 2026'           -> 'April 2026'   (unchanged)
        'April 2026 20260520'  -> 'April 2026'   (fixed)
        'APRIL 2026 20260520'  -> 'APRIL 2026'   (fixed)
    """
    parts = period_label.strip().split()
    if len(parts) >= 2:
        return f"{parts[0]} {parts[1]}"
    return period_label.strip()


# ── Formatting helpers ────────────────────────────────────────────────────────

def _fmt(value, currency=True, currency_symbol=REPORT_CURRENCY):
    try:
        formatted = f"{int(value):,}"
        if currency:
            return f"{currency_symbol} {formatted}"
        return formatted
    except (TypeError, ValueError):
        return str(value)


def _pct(value):
    try:
        return f"{float(value):.1f}%"
    except (TypeError, ValueError):
        return str(value)


def _safe(metrics, key, default="N/A"):
    return metrics.get(key, default)


def _filename(prefix, period_label, extension):
    clean = _clean_period(period_label).replace(" ", "_")
    return f"{prefix}_{clean}.{extension}"


def _xl_fill(hex_color):
    return PatternFill(fill_type="solid", fgColor=hex_color)


def _xl_border(color="CCCCCC"):
    side = Side(style="thin", color=color)
    return Border(left=side, right=side, top=side, bottom=side)


def _xl_header_font(color="FFFFFF", size=10, bold=True):
    return Font(name="Calibri", bold=bold, size=size, color=color)


def _xl_body_font(size=10, bold=False, color="222222"):
    return Font(name="Calibri", size=size, bold=bold, color=color)


def _xl_write_header_row(ws, row, values, fill_color, font_color="FFFFFF", height=20):
    ws.row_dimensions[row].height = height
    for col, value in enumerate(values, start=1):
        cell = ws.cell(row=row, column=col, value=value)
        cell.fill      = _xl_fill(fill_color)
        cell.font      = _xl_header_font(color=font_color)
        cell.border    = _xl_border()
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)


def _xl_write_data_row(ws, row, values, alt=False, bold=False, number_cols=None):
    fill = _xl_fill(XL_LIGHT) if alt else _xl_fill(XL_WHITE)
    number_cols = number_cols or []
    ws.row_dimensions[row].height = 18
    for col, value in enumerate(values, start=1):
        cell = ws.cell(row=row, column=col, value=value)
        cell.fill      = fill
        cell.font      = _xl_body_font(bold=bold)
        cell.border    = _xl_border()
        align          = "right" if col in number_cols else "left"
        cell.alignment = Alignment(horizontal=align, vertical="center")


# =============================================================
# EXPORT ENGINE
# =============================================================

class ExportEngine:

    def __init__(self, station_config: dict = None):
        """
        Initialise the export engine.

        Args:
            station_config: Dict from load_station_config(). When provided,
                            output directories and station identity come from it.
                            When None, falls back to config.settings (legacy mode).
        """
        self.station_config = station_config or {}

        # ── Resolve output directories ────────────────────────────────────────
        if station_config and "report_dirs" in station_config:
            self.pdf_dir  = station_config["report_dirs"]["pdf"]
            self.docx_dir = station_config["report_dirs"]["docx"]
            self.xlsx_dir = station_config["report_dirs"]["xlsx"]
        else:
            # Legacy fallback — uses hardcoded dirs from config.settings
            self.pdf_dir  = REPORTS_PDF_DIR
            self.docx_dir = REPORTS_DOCX_DIR
            self.xlsx_dir = REPORTS_PDF_DIR.parent / "xlsx"

        # ── Resolve station identity ──────────────────────────────────────────
        if station_config:
            self.station_name     = station_config.get("station_name", STATION_NAME)
            self.station_location = station_config.get("location", STATION_LOCATION)
            self.currency         = station_config.get("currency", REPORT_CURRENCY)
        else:
            self.station_name     = STATION_NAME
            self.station_location = STATION_LOCATION
            self.currency         = REPORT_CURRENCY

        # ── Create output directories if they don't exist ─────────────────────
        self.pdf_dir.mkdir(parents=True, exist_ok=True)
        self.docx_dir.mkdir(parents=True, exist_ok=True)
        self.xlsx_dir.mkdir(parents=True, exist_ok=True)

    # =========================================================
    # PDF
    # =========================================================

    def generate_pdf(self, metrics, report_sections, period_label):
        # ── Sanitise period label before anything uses it ─────────────────────
        period_label = _clean_period(period_label)

        filename = _filename("StationDeck_Report", period_label, "pdf")
        filepath = self.pdf_dir / filename

        doc = SimpleDocTemplate(
            str(filepath), pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )
        styles = self._pdf_styles()
        story  = []
        story += self._pdf_cover(styles, period_label)
        story.append(PageBreak())
        story += self._pdf_section(styles, "Executive Summary",
                                   report_sections.get("executive_summary", ""))
        story += self._pdf_fuel_table(styles, metrics)
        story += self._pdf_section(styles, "Fuel Sales Analysis",
                                   report_sections.get("fuel_sales_analysis", ""))
        story += self._pdf_payment_table(styles, metrics)
        story += self._pdf_section(styles, "Payment Collection Analysis",
                                   report_sections.get("payment_collection_analysis", ""))
        story += self._pdf_expense_table(styles, metrics)
        story += self._pdf_reconciliation_table(styles, metrics)
        story += self._pdf_section(styles, "Cash Reconciliation Analysis",
                                   report_sections.get("cash_reconciliation_analysis", ""))
        story.append(Spacer(1, 0.5*cm))
        story.append(HRFlowable(width="100%", thickness=1, color=BRAND_MID))
        story.append(Paragraph(
            f"Generated by {REPORT_AUTHOR} \u00b7 "
            f"{datetime.now().strftime('%d %B %Y, %H:%M')}",
            styles["footer"]
        ))
        doc.build(story)
        print(f"  OK  PDF saved: {filepath}")
        return str(filepath)

    def _pdf_styles(self):
        styles = {}
        styles["cover_station"]   = ParagraphStyle("cover_station",  fontName="Helvetica-Bold", fontSize=26, textColor=white, alignment=TA_CENTER, spaceAfter=6)
        styles["cover_location"]  = ParagraphStyle("cover_location", fontName="Helvetica", fontSize=13, textColor=HexColor("#DDDDDD"), alignment=TA_CENTER, spaceAfter=4)
        styles["cover_title"]     = ParagraphStyle("cover_title",    fontName="Helvetica-Bold", fontSize=16, textColor=white, alignment=TA_CENTER, spaceAfter=4)
        styles["cover_period"]    = ParagraphStyle("cover_period",   fontName="Helvetica-Bold", fontSize=22, textColor=white, alignment=TA_CENTER, spaceAfter=8)
        styles["cover_generated"] = ParagraphStyle("cover_generated",fontName="Helvetica", fontSize=10, textColor=HexColor("#AAAAAA"), alignment=TA_CENTER)
        styles["section_heading"] = ParagraphStyle("section_heading",fontName="Helvetica-Bold", fontSize=13, textColor=BRAND_RED, spaceBefore=14, spaceAfter=6)
        styles["body"]            = ParagraphStyle("body",           fontName="Helvetica", fontSize=10, textColor=TEXT_DARK, leading=15, spaceAfter=8)
        styles["table_header"]    = ParagraphStyle("table_header",   fontName="Helvetica-Bold", fontSize=9, textColor=white, alignment=TA_LEFT)
        styles["table_cell"]      = ParagraphStyle("table_cell",     fontName="Helvetica", fontSize=9, textColor=TEXT_DARK)
        styles["table_cell_right"]= ParagraphStyle("table_cell_right",fontName="Helvetica", fontSize=9, textColor=TEXT_DARK, alignment=TA_RIGHT)
        styles["footer"]          = ParagraphStyle("footer",         fontName="Helvetica", fontSize=8, textColor=HexColor("#888888"), alignment=TA_CENTER, spaceBefore=4)
        # Logo styles (used in cover)
        styles["logo_station"]    = ParagraphStyle("logo_station",   fontName="Helvetica-Bold", fontSize=18, textColor=HexColor("#1A1A2E"), alignment=TA_CENTER)
        styles["logo_deck"]       = ParagraphStyle("logo_deck",      fontName="Helvetica-Bold", fontSize=18, textColor=BRAND_RED, alignment=TA_CENTER)
        styles["logo_tagline"]    = ParagraphStyle("logo_tagline",   fontName="Helvetica", fontSize=8, textColor=HexColor("#9CA3AF"), alignment=TA_CENTER, spaceBefore=2)
        return styles

    def _pdf_cover(self, styles, period_label):
        """
        Cover page layout:
          [spacer]
          [dark navy block: station name, location, report title, period, generated by]
          [red divider]
          [StationDeck logo: red dot + Station(dark)Deck(red) wordmark + tagline]
        """
        elements = []

        # ── Station identity block ─────────────────────────────────────────────
        cover_data = [
            [Paragraph(self.station_name, styles["cover_station"])],
            [Paragraph(self.station_location, styles["cover_location"])],
            [Spacer(1, 1*cm)],
            [Paragraph("OPERATIONAL PERFORMANCE REPORT", styles["cover_title"])],
            [Paragraph(period_label.upper(), styles["cover_period"])],
            [Spacer(1, 0.5*cm)],
            [Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}", styles["cover_generated"])],
            [Paragraph(f"Prepared by: {REPORT_AUTHOR}", styles["cover_generated"])],
        ]
        ct = Table(cover_data, colWidths=[17*cm])
        ct.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,-1),BRAND_DARK),
            ("TOPPADDING",(0,0),(-1,-1),12),
            ("BOTTOMPADDING",(0,0),(-1,-1),12),
            ("LEFTPADDING",(0,0),(-1,-1),20),
            ("RIGHTPADDING",(0,0),(-1,-1),20),
            ("ROUNDEDCORNERS",[8]),
        ]))
        elements.append(Spacer(1, 3*cm))
        elements.append(ct)
        elements.append(Spacer(1, 1*cm))

        # ── Red divider ────────────────────────────────────────────────────────
        elements.append(HRFlowable(width="100%", thickness=4, color=BRAND_RED))
        elements.append(Spacer(1, 0.8*cm))

        # ── StationDeck logo (Variation 1: light) ─────────────────────────────
        # Built from a single-row table so the dot + wordmark sit on one baseline.
        #
        # Layout: [red dot cell] [Station (dark)] [Deck (red)]
        # Then tagline centered below.
        #
        # We use a 3-column table for the inline row, then wrap in an outer
        # centered table so the whole group is centered on the page.

        dot_style = ParagraphStyle(
            "logo_dot_pdf",
            fontName="Helvetica-Bold",
            fontSize=22,
            textColor=BRAND_RED,
            alignment=TA_CENTER,
            leading=24,
        )
        station_style = ParagraphStyle(
            "logo_word_station",
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=HexColor("#1A1A2E"),
            alignment=TA_LEFT,
            leading=24,
        )
        deck_style = ParagraphStyle(
            "logo_word_deck",
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=BRAND_RED,
            alignment=TA_LEFT,
            leading=24,
        )
        tagline_style = ParagraphStyle(
            "logo_tagline_pdf",
            fontName="Helvetica",
            fontSize=8,
            textColor=HexColor("#9CA3AF"),
            alignment=TA_CENTER,
            spaceBefore=4,
            letterSpacing=1.5,
        )

        wordmark_row = Table(
            [[
                Paragraph("\u25cf", dot_style),       # ● red dot
                Paragraph("Station", station_style),
                Paragraph("Deck", deck_style),
            ]],
            colWidths=[0.6*cm, 3.2*cm, 2.2*cm],
        )
        wordmark_row.setStyle(TableStyle([
            ("ALIGN",   (0,0), (-1,-1), "CENTER"),
            ("VALIGN",  (0,0), (-1,-1), "MIDDLE"),
            ("LEFTPADDING",  (0,0), (-1,-1), 2),
            ("RIGHTPADDING", (0,0), (-1,-1), 2),
            ("TOPPADDING",   (0,0), (-1,-1), 0),
            ("BOTTOMPADDING",(0,0), (-1,-1), 0),
        ]))

        tagline_para = Paragraph("FUEL STATION REPORTING SYSTEM", tagline_style)

        logo_block = Table(
            [[wordmark_row], [tagline_para]],
            colWidths=[8*cm],
        )
        logo_block.setStyle(TableStyle([
            ("ALIGN",   (0,0), (-1,-1), "CENTER"),
            ("VALIGN",  (0,0), (-1,-1), "MIDDLE"),
            ("TOPPADDING",   (0,0), (-1,-1), 4),
            ("BOTTOMPADDING",(0,0), (-1,-1), 4),
        ]))

        # Center the logo block on the page using an outer wrapper table
        outer = Table([[logo_block]], colWidths=[17*cm])
        outer.setStyle(TableStyle([
            ("ALIGN",  (0,0), (-1,-1), "CENTER"),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ]))

        elements.append(outer)
        return elements

    def _pdf_section(self, styles, title, body):
        elements = []
        elements.append(Paragraph(title.upper(), styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        if body:
            for para in body.strip().split("\n\n"):
                elements.append(Paragraph(para.strip(), styles["body"]))
        return elements

    def _pdf_fuel_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("FUEL PERFORMANCE SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        data = [
            [Paragraph("Metric", H), Paragraph("PMS (Petrol)", H), Paragraph("AGO (Diesel)", H)],
            [Paragraph("Volume Sold", C),
             Paragraph(_fmt(_safe(metrics,"pms_volume_total"),False)+" L", R),
             Paragraph(_fmt(_safe(metrics,"ago_volume_total"),False)+" L", R)],
            [Paragraph("Revenue", C),
             Paragraph(_fmt(_safe(metrics,"pms_revenue_total"), currency_symbol=cur), R),
             Paragraph(_fmt(_safe(metrics,"ago_revenue_total"), currency_symbol=cur), R)],
            [Paragraph("Combined Revenue", C), Paragraph("", C),
             Paragraph(_fmt(_safe(metrics,"total_fuel_revenue"), currency_symbol=cur), R)],
        ]
        t = Table(data, colWidths=[6*cm, 5.5*cm, 5.5*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_payment_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("PAYMENT COLLECTION BREAKDOWN", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        data = [
            [Paragraph("Payment Type", H), Paragraph("Amount", H), Paragraph("Share", H)],
            [Paragraph("Cash Collected", C),     Paragraph(_fmt(_safe(metrics,"cash_collected"), currency_symbol=cur), R),    Paragraph(_pct(_safe(metrics,"cash_percentage")), R)],
            [Paragraph("Cashless Total", C),      Paragraph(_fmt(_safe(metrics,"cashless_collected"), currency_symbol=cur), R), Paragraph(_pct(_safe(metrics,"cashless_percentage")), R)],
            [Paragraph("  \u21b3 Plus Card", C),  Paragraph(_fmt(_safe(metrics,"plus_card_total"), currency_symbol=cur), R),   Paragraph("", C)],
            [Paragraph("  \u21b3 Visa", C),       Paragraph(_fmt(_safe(metrics,"visa_total"), currency_symbol=cur), R),         Paragraph("", C)],
            [Paragraph("  \u21b3 Credit Sales", C),Paragraph(_fmt(_safe(metrics,"credit_sales_total"), currency_symbol=cur), R),Paragraph("", C)],
            [Paragraph("Total Sales", C),          Paragraph(_fmt(_safe(metrics,"total_sales"), currency_symbol=cur), R),       Paragraph("100%", R)],
        ]
        t = Table(data, colWidths=[6*cm, 6*cm, 5*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_expense_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("EXPENSE SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        data = [
            [Paragraph("Item", H), Paragraph("Amount", H)],
            [Paragraph("Total Expenses", C),        Paragraph(_fmt(_safe(metrics,"total_expenses"), currency_symbol=cur), R)],
            [Paragraph("Total Revenue", C),          Paragraph(_fmt(_safe(metrics,"total_revenue"), currency_symbol=cur), R)],
            [Paragraph("Avg Daily Fuel Revenue", C), Paragraph(_fmt(_safe(metrics,"avg_daily_fuel_revenue"), currency_symbol=cur), R)],
        ]
        t = Table(data, colWidths=[9*cm, 8*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_reconciliation_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("CASH RECONCILIATION SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        delta_status = str(_safe(metrics, "delta_status", "UNKNOWN"))
        delta_color  = SURPLUS_GREEN if delta_status == "SURPLUS" else DEFICIT_RED
        data = [
            [Paragraph("Item", H), Paragraph("Value", H)],
            [Paragraph("Cash Collected", C),      Paragraph(_fmt(_safe(metrics,"cash_collected"), currency_symbol=cur), R)],
            [Paragraph("Total Banked", C),         Paragraph(_fmt(_safe(metrics,"total_cash_banked"), currency_symbol=cur), R)],
            [Paragraph("Expected to Bank", C),     Paragraph(_fmt(_safe(metrics,"total_cash_expected"), currency_symbol=cur), R)],
            [Paragraph("Net Delta", C),            Paragraph(_fmt(_safe(metrics,"total_delta"), currency_symbol=cur), R)],
            [Paragraph("Delta Status", C),         Paragraph(delta_status, R)],
            [Paragraph("Days with Anomalies", C),  Paragraph(str(_safe(metrics,"anomaly_days_count")), R)],
        ]
        t = Table(data, colWidths=[9*cm, 8*cm])
        style = self._base_table_style()
        t.setStyle(TableStyle([
            *style._cmds,
            ("TEXTCOLOR",(1,5),(1,5), delta_color),
            ("FONTNAME",(1,5),(1,5), "Helvetica-Bold"),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _base_table_style(self):
        return TableStyle([
            ("BACKGROUND",(0,0),(-1,0),BRAND_DARK),
            ("TEXTCOLOR",(0,0),(-1,0),white),
            ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
            ("FONTSIZE",(0,0),(-1,0),9),
            ("TOPPADDING",(0,0),(-1,0),7),
            ("BOTTOMPADDING",(0,0),(-1,0),7),
            ("FONTNAME",(0,1),(-1,-1),"Helvetica"),
            ("FONTSIZE",(0,1),(-1,-1),9),
            ("TOPPADDING",(0,1),(-1,-1),5),
            ("BOTTOMPADDING",(0,1),(-1,-1),5),
            ("LEFTPADDING",(0,0),(-1,-1),8),
            ("RIGHTPADDING",(0,0),(-1,-1),8),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[white, BRAND_LIGHT]),
            ("GRID",(0,0),(-1,-1),0.5,BRAND_MID),
            ("LINEABOVE",(0,0),(-1,0),1.5,BRAND_RED),
        ])

    # =========================================================
    # DOCX
    # =========================================================

    def generate_docx(self, metrics, report_sections, period_label):
        # ── Sanitise period label before anything uses it ─────────────────────
        period_label = _clean_period(period_label)

        filename = _filename("StationDeck_Report", period_label, "docx")
        filepath = self.docx_dir / filename

        doc = Document()
        self._docx_set_margins(doc)
        self._docx_cover(doc, period_label)
        doc.add_page_break()
        self._docx_section_heading(doc, "Executive Summary")
        self._docx_body(doc, report_sections.get("executive_summary", ""))
        self._docx_section_heading(doc, "Fuel Performance Summary")
        self._docx_fuel_table(doc, metrics)
        self._docx_body(doc, report_sections.get("fuel_sales_analysis", ""))
        self._docx_section_heading(doc, "Payment Collection Breakdown")
        self._docx_payment_table(doc, metrics)
        self._docx_body(doc, report_sections.get("payment_collection_analysis", ""))
        self._docx_section_heading(doc, "Expense Summary")
        self._docx_expense_table(doc, metrics)
        self._docx_section_heading(doc, "Cash Reconciliation Summary")
        self._docx_reconciliation_table(doc, metrics)
        self._docx_body(doc, report_sections.get("cash_reconciliation_analysis", ""))
        self._docx_footer_line(doc, period_label)
        doc.save(str(filepath))
        print(f"  OK  DOCX saved: {filepath}")
        return str(filepath)

    def _docx_set_margins(self, doc):
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2)
            section.right_margin  = Cm(2)

    def _docx_cover(self, doc, period_label):
        """
        Cover page layout:
          Station name (large, dark)
          Location (grey)
          [blank line]
          OPERATIONAL PERFORMANCE REPORT (red)
          MONTH YEAR (large, dark)
          [blank line]
          Generated / Prepared by (grey)
          [blank line]
          ─────────── red divider ───────────
          [blank line]
          ● StationDeck  (logo: dot + Station dark + Deck red)
          FUEL STATION REPORTING SYSTEM  (tagline, grey)
        """
        # Station name
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.station_name)
        run.font.size = Pt(28); run.font.bold = True; run.font.color.rgb = W_DARK

        # Location
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.station_location)
        run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        # Report type
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("OPERATIONAL PERFORMANCE REPORT")
        run.font.size = Pt(14); run.font.bold = True; run.font.color.rgb = W_RED

        # Period — CLEAN label, no date stamp
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(period_label.upper())
        run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = W_DARK

        doc.add_paragraph()

        # Generated / prepared by
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
        run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Prepared by: {REPORT_AUTHOR}")
        run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        # ── Red divider line ───────────────────────────────────────────────────
        # Achieved via a bottom border on an empty paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "24")       # thickness in eighths of a point (24 = 3pt)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "C8102E")   # brand red
        pBdr.append(bottom)
        pPr.append(pBdr)

        doc.add_paragraph()

        # ── StationDeck logo (Variation 1) ─────────────────────────────────────
        # Single paragraph with mixed runs: ● (red) + Station (dark) + Deck (red)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)

        # Red dot
        run = p.add_run("\u25cf ")          # ●
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = W_RED

        # "Station" in dark navy
        run = p.add_run("Station")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = W_DARK

        # "Deck" in red
        run = p.add_run("Deck")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = W_RED

        # Tagline on next line
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run("FUEL STATION REPORTING SYSTEM")
        run.font.size = Pt(7)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    def _docx_section_heading(self, doc, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(title.upper())
        run.font.size = Pt(12); run.font.bold = True; run.font.color.rgb = W_RED

    def _docx_body(self, doc, text):
        if not text:
            return
        for para in text.strip().split("\n\n"):
            p = doc.add_paragraph()
            run = p.add_run(para.strip())
            run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            p.paragraph_format.space_after = Pt(6)

    def _docx_fuel_table(self, doc, metrics):
        cur = self.currency
        self._docx_table(doc, ["Metric", "PMS (Petrol)", "AGO (Diesel)"], [
            ["Volume Sold",
             _fmt(_safe(metrics,"pms_volume_total"), False)+" L",
             _fmt(_safe(metrics,"ago_volume_total"), False)+" L"],
            ["Revenue",
             _fmt(_safe(metrics,"pms_revenue_total"), currency_symbol=cur),
             _fmt(_safe(metrics,"ago_revenue_total"), currency_symbol=cur)],
            ["Combined Revenue", "",
             _fmt(_safe(metrics,"total_fuel_revenue"), currency_symbol=cur)],
        ])

    def _docx_payment_table(self, doc, metrics):
        cur = self.currency
        self._docx_table(doc, ["Payment Type", "Amount", "Share"], [
            ["Cash Collected",      _fmt(_safe(metrics,"cash_collected"), currency_symbol=cur),      _pct(_safe(metrics,"cash_percentage"))],
            ["Cashless Total",       _fmt(_safe(metrics,"cashless_collected"), currency_symbol=cur),  _pct(_safe(metrics,"cashless_percentage"))],
            ["  \u21b3 Plus Card",  _fmt(_safe(metrics,"plus_card_total"), currency_symbol=cur),     ""],
            ["  \u21b3 Visa",       _fmt(_safe(metrics,"visa_total"), currency_symbol=cur),           ""],
            ["  \u21b3 Credit Sales",_fmt(_safe(metrics,"credit_sales_total"), currency_symbol=cur), ""],
            ["Total Sales",          _fmt(_safe(metrics,"total_sales"), currency_symbol=cur),         "100%"],
        ])

    def _docx_expense_table(self, doc, metrics):
        cur = self.currency
        self._docx_table(doc, ["Item", "Amount"], [
            ["Total Expenses",        _fmt(_safe(metrics,"total_expenses"), currency_symbol=cur)],
            ["Total Revenue",          _fmt(_safe(metrics,"total_revenue"), currency_symbol=cur)],
            ["Avg Daily Fuel Revenue", _fmt(_safe(metrics,"avg_daily_fuel_revenue"), currency_symbol=cur)],
        ])

    def _docx_reconciliation_table(self, doc, metrics):
        cur = self.currency
        self._docx_table(doc, ["Item", "Value"], [
            ["Cash Collected",      _fmt(_safe(metrics,"cash_collected"), currency_symbol=cur)],
            ["Total Banked",         _fmt(_safe(metrics,"total_cash_banked"), currency_symbol=cur)],
            ["Expected to Bank",     _fmt(_safe(metrics,"total_cash_expected"), currency_symbol=cur)],
            ["Net Delta",            _fmt(_safe(metrics,"total_delta"), currency_symbol=cur)],
            ["Delta Status",         str(_safe(metrics,"delta_status","UNKNOWN"))],
            ["Days with Anomalies",  str(_safe(metrics,"anomaly_days_count"))],
        ])

    def _docx_table(self, doc, headers, rows):
        table = doc.add_table(rows=1+len(rows), cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(table.rows[0].cells):
            h.text = ""
            run = h.paragraphs[0].add_run(headers[i])
            run.font.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            self._docx_set_cell_bg(h, "1A1A2E")
        for r_idx, row_data in enumerate(rows):
            bg = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, cell in enumerate(table.rows[r_idx+1].cells):
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(row_data[c_idx]))
                run.font.size = Pt(9)
                self._docx_set_cell_bg(cell, bg)
        doc.add_paragraph()

    def _docx_set_cell_bg(self, cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _docx_footer_line(self, doc, period_label):
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"StationDeck Report \u00b7 {period_label} \u00b7 "
            f"Generated {datetime.now().strftime('%d %B %Y, %H:%M')} \u00b7 {REPORT_AUTHOR}"
        )
        run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # =========================================================
    # EXCEL
    # =========================================================

    def generate_xlsx(self, metrics, daily_df, period_label):
        # ── Sanitise period label before anything uses it ─────────────────────
        period_label = _clean_period(period_label)

        filename = _filename("StationDeck_Report", period_label, "xlsx")
        filepath = self.xlsx_dir / filename

        wb = openpyxl.Workbook()
        self._xlsx_summary_sheet(wb, metrics, period_label)
        self._xlsx_daily_sheet(wb, daily_df, period_label)
        if "Sheet" in wb.sheetnames:
            del wb["Sheet"]
        wb.save(str(filepath))
        print(f"  OK  XLSX saved: {filepath}")
        return str(filepath)

    def _xlsx_summary_sheet(self, wb, metrics, period_label):
        ws = wb.create_sheet(title="Monthly Summary")
        ws.sheet_view.showGridLines = False
        cur = self.currency

        # Row 1: Station name — period label is already clean here
        ws.merge_cells("A1:C1")
        c = ws["A1"]
        c.value = f"{self.station_name} \u2014 {self.station_location}"
        c.font      = Font(name="Calibri", bold=True, size=14, color="FFFFFF")
        c.fill      = _xl_fill(XL_DARK)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 28

        # Row 2: Report title — period_label is clean, no date stamp
        ws.merge_cells("A2:C2")
        c = ws["A2"]
        c.value = f"Operational Performance Report \u2014 {period_label}"
        c.font      = Font(name="Calibri", size=11, color="FFFFFF")
        c.fill      = _xl_fill(XL_SUBHEAD)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[2].height = 22

        # Row 3: Generated by
        ws.merge_cells("A3:C3")
        c = ws["A3"]
        c.value = f"Generated by {REPORT_AUTHOR} on {datetime.now().strftime('%d %B %Y, %H:%M')}"
        c.font      = Font(name="Calibri", size=9, color="888888")
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[3].height = 16

        def write_section(start_row, section_title, rows_data):
            ws.row_dimensions[start_row].height = 8
            start_row += 1
            ws.merge_cells(f"A{start_row}:C{start_row}")
            hdr = ws[f"A{start_row}"]
            hdr.value     = section_title
            hdr.font      = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
            hdr.fill      = _xl_fill(XL_RED)
            hdr.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            ws.row_dimensions[start_row].height = 18
            start_row += 1
            _xl_write_header_row(ws, start_row, ["Metric", "Value", ""], XL_DARK, height=16)
            start_row += 1
            for i, (label, value) in enumerate(rows_data):
                _xl_write_data_row(ws, start_row, [label, value, ""],
                                   alt=(i % 2 == 0), number_cols=[2])
                start_row += 1
            return start_row

        row = write_section(4, "FUEL PERFORMANCE", [
            ("PMS Volume Sold",      f"{_safe(metrics,'pms_volume_total',0):,.2f} L"),
            ("AGO Volume Sold",      f"{_safe(metrics,'ago_volume_total',0):,.2f} L"),
            ("PMS Revenue",          _fmt(_safe(metrics,"pms_revenue_total"), currency_symbol=cur)),
            ("AGO Revenue",          _fmt(_safe(metrics,"ago_revenue_total"), currency_symbol=cur)),
            ("Total Fuel Revenue",   _fmt(_safe(metrics,"total_fuel_revenue"), currency_symbol=cur)),
            ("Avg Daily Fuel Rev.",  _fmt(_safe(metrics,"avg_daily_fuel_revenue"), currency_symbol=cur)),
        ])
        row = write_section(row, "NON-FUEL REVENUE", [
            ("Lubricants Revenue",   _fmt(_safe(metrics,"lubes_revenue_total"), currency_symbol=cur)),
            ("LPG Gas Revenue",      _fmt(_safe(metrics,"lpg_revenue_total"), currency_symbol=cur)),
            ("Shop Sales",           _fmt(_safe(metrics,"shop_sales_total"), currency_symbol=cur)),
            ("Total Non-Fuel Rev.",  _fmt(_safe(metrics,"total_nonfuel_revenue"), currency_symbol=cur)),
            ("Total Sales",          _fmt(_safe(metrics,"total_sales"), currency_symbol=cur)),
            ("Total Revenue",        _fmt(_safe(metrics,"total_revenue"), currency_symbol=cur)),
        ])
        row = write_section(row, "PAYMENT COLLECTION", [
            ("Cash Collected",       _fmt(_safe(metrics,"cash_collected"), currency_symbol=cur)),
            ("Cash Share",           _pct(_safe(metrics,"cash_percentage"))),
            ("Cashless Collected",   _fmt(_safe(metrics,"cashless_collected"), currency_symbol=cur)),
            ("Cashless Share",       _pct(_safe(metrics,"cashless_percentage"))),
            ("Plus Card Total",      _fmt(_safe(metrics,"plus_card_total"), currency_symbol=cur)),
            ("Visa Total",           _fmt(_safe(metrics,"visa_total"), currency_symbol=cur)),
            ("Credit Sales",         _fmt(_safe(metrics,"credit_sales_total"), currency_symbol=cur)),
        ])
        row = write_section(row, "EXPENSES", [
            ("Total Expenses",       _fmt(_safe(metrics,"total_expenses"), currency_symbol=cur)),
        ])

        delta_status = str(_safe(metrics, "delta_status", "UNKNOWN"))
        delta_color  = XL_GREEN if delta_status == "SURPLUS" else XL_RED

        row = write_section(row, "CASH RECONCILIATION", [
            ("Days Covered",         str(_safe(metrics,"total_days"))),
            ("Cash Collected",       _fmt(_safe(metrics,"cash_collected"), currency_symbol=cur)),
            ("Total Banked",         _fmt(_safe(metrics,"total_cash_banked"), currency_symbol=cur)),
            ("Expected to Bank",     _fmt(_safe(metrics,"total_cash_expected"), currency_symbol=cur)),
            ("Net Delta",            _fmt(_safe(metrics,"total_delta"), currency_symbol=cur)),
            ("Delta Status",         delta_status),
            ("Days with Anomalies",  str(_safe(metrics,"anomaly_days_count"))),
        ])

        status_cell = ws[f"B{row-2}"]
        status_cell.font = Font(name="Calibri", bold=True, size=10, color=delta_color)

        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 22
        ws.column_dimensions["C"].width = 4

    def _xlsx_daily_sheet(self, wb, daily_df, period_label):
        if daily_df is None or daily_df.empty:
            return

        ws = wb.create_sheet(title="Daily Records")
        ws.sheet_view.showGridLines = False
        col_count = 10

        ws.merge_cells(f"A1:{get_column_letter(col_count)}1")
        c = ws["A1"]
        c.value = f"{self.station_name} \u2014 Daily Records \u2014 {period_label}"
        c.font      = Font(name="Calibri", bold=True, size=13, color="FFFFFF")
        c.fill      = _xl_fill(XL_DARK)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 26

        _xl_write_header_row(ws, 2, [
            "Date","PMS Vol (L)","AGO Vol (L)","PMS Rev (UGX)","AGO Rev (UGX)",
            "Total Sales (UGX)","Cash (UGX)","Cashless (UGX)","Expenses (UGX)","Delta (UGX)"
        ], XL_DARK, height=20)

        number_cols = list(range(2, col_count+1))

        for i, (_, row_data) in enumerate(daily_df.iterrows()):
            excel_row = i + 3
            alt = (i % 2 == 0)

            def safe_val(col_name, decimals=0, _r=row_data):
                try:
                    val = _r.get(col_name, 0)
                    return round(float(val), decimals) if decimals else int(float(val))
                except:
                    return 0

            try:
                dv = row_data["date"]
                date_str = dv.strftime("%d %b %Y") if hasattr(dv, "strftime") else str(dv)
            except:
                date_str = "N/A"

            cashless_val = safe_val("cashless_total") - safe_val("total_cash")
            values = [
                date_str,
                safe_val("pms_volume", 2),
                safe_val("ago_volume", 2),
                safe_val("pms_revenue"),
                safe_val("ago_revenue"),
                safe_val("cashless_total"),
                safe_val("total_cash"),
                cashless_val,
                safe_val("total_expenses"),
                safe_val("delta"),
            ]
            _xl_write_data_row(ws, excel_row, values, alt=alt, number_cols=number_cols)

            dv = values[9]
            if isinstance(dv, (int, float)):
                ws[f"J{excel_row}"].font = Font(
                    name="Calibri", size=10,
                    color=XL_GREEN if dv >= 0 else XL_RED
                )

        for i, w in enumerate([14,13,13,18,18,18,16,16,16,14], start=1):
            ws.column_dimensions[get_column_letter(i)].width = w

        ws.freeze_panes = "A3"