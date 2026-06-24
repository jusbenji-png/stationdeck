# =============================================================
# StationDeck - Export Engine
# src/exporter.py
# =============================================================
# Generates PDF, DOCX, and XLSX report files.
#
# Monthly usage:
#   engine = ExportEngine(station_config=station_config)
#   pdf_path  = engine.generate_pdf(metrics, sections, period_label)
#   docx_path = engine.generate_docx(metrics, sections, period_label)
#   xlsx_path = engine.generate_xlsx(metrics, daily_df, period_label)
#
# Annual usage:
#   pdf_path  = engine.generate_annual_pdf(metrics, sections, period_label)
#   docx_path = engine.generate_annual_docx(metrics, sections, period_label)
#   xlsx_path = engine.generate_annual_xlsx(metrics, period_label)
# =============================================================

from pathlib import Path
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.colors import HexColor, white
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from config.settings import (
    REPORTS_PDF_DIR, REPORTS_DOCX_DIR,
    STATION_NAME, STATION_LOCATION,
    REPORT_CURRENCY, REPORT_AUTHOR,
)

# ── Colour constants ──────────────────────────────────────────────────────────
BRAND_RED     = HexColor("#C8102E")
BRAND_DARK    = HexColor("#1A1A2E")
BRAND_LIGHT   = HexColor("#F5F5F5")
BRAND_MID     = HexColor("#E0E0E0")
TEXT_DARK     = HexColor("#222222")
SURPLUS_GREEN = HexColor("#1E7E34")
DEFICIT_RED   = HexColor("#C8102E")

W_RED  = RGBColor(0xC8, 0x10, 0x2E)
W_DARK = RGBColor(0x1A, 0x1A, 0x2E)
W_GREY = RGBColor(0xF5, 0xF5, 0xF5)
W_MID  = RGBColor(0xE0, 0xE0, 0xE0)

XL_DARK    = "1A1A2E"
XL_RED     = "C8102E"
XL_LIGHT   = "F5F5F5"
XL_WHITE   = "FFFFFF"
XL_MID     = "E0E0E0"
XL_GREEN   = "1E7E34"
XL_SUBHEAD = "2E4057"
XL_ORANGE  = "E8660A"


# ── Period label sanitiser ────────────────────────────────────────────────────
def _clean_period(period_label):
    """Keep only first two tokens — guards against date-stamp concatenation."""
    parts = period_label.strip().split()
    if len(parts) >= 2:
        return f"{parts[0]} {parts[1]}"
    return period_label.strip()


# ── Formatting helpers ────────────────────────────────────────────────────────
def _fmt(value, currency=True, currency_symbol=REPORT_CURRENCY):
    try:
        formatted = f"{int(value):,}"
        return f"{currency_symbol} {formatted}" if currency else formatted
    except (TypeError, ValueError):
        return str(value)

def _pct(value):
    try:    return f"{float(value):.1f}%"
    except: return str(value)

def _vol(value):
    try:    return f"{float(value):,.2f} L"
    except: return str(value)

def _safe(metrics, key, default="N/A"):
    return metrics.get(key, default)

def _filename(prefix, period_label, extension):
    clean = _clean_period(period_label).replace(" ", "_")
    return f"{prefix}_{clean}.{extension}"

def _annual_filename(prefix, fy_label, extension):
    clean = fy_label.replace("/", "-").replace(" ", "_")
    return f"{prefix}_Annual_{clean}.{extension}"


# ── Excel helpers ─────────────────────────────────────────────────────────────
def _xl_fill(hex_color):
    return PatternFill(fill_type="solid", fgColor=hex_color)

def _xl_border(color="CCCCCC"):
    side = Side(style="thin", color=color)
    return Border(left=side, right=side, top=side, bottom=side)

def _xl_header_font(color="FFFFFF", size=10, bold=True):
    return Font(name="Calibri", bold=bold, size=size, color=color)

def _xl_body_font(size=10, bold=False, color="222222"):
    return Font(name="Calibri", size=size, bold=bold, color=color)

def _xl_write_header_row(ws, row, values, fill_color, font_color="FFFFFF", height=20):
    ws.row_dimensions[row].height = height
    for col, value in enumerate(values, start=1):
        cell = ws.cell(row=row, column=col, value=value)
        cell.fill      = _xl_fill(fill_color)
        cell.font      = _xl_header_font(color=font_color)
        cell.border    = _xl_border()
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

def _xl_write_data_row(ws, row, values, alt=False, bold=False, number_cols=None):
    fill = _xl_fill(XL_LIGHT) if alt else _xl_fill(XL_WHITE)
    number_cols = number_cols or []
    ws.row_dimensions[row].height = 18
    for col, value in enumerate(values, start=1):
        cell = ws.cell(row=row, column=col, value=value)
        cell.fill      = fill
        cell.font      = _xl_body_font(bold=bold)
        cell.border    = _xl_border()
        align          = "right" if col in number_cols else "left"
        cell.alignment = Alignment(horizontal=align, vertical="center")

def _xl_sheet_header(ws, title, subtitle, col_count, cur_row=1):
    """Write a standard 3-row sheet header. Returns next available row."""
    ws.sheet_view.showGridLines = False
    last_col = get_column_letter(col_count)

    ws.merge_cells(f"A{cur_row}:{last_col}{cur_row}")
    c = ws[f"A{cur_row}"]
    c.value = title
    c.font      = Font(name="Calibri", bold=True, size=13, color="FFFFFF")
    c.fill      = _xl_fill(XL_DARK)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[cur_row].height = 26
    cur_row += 1

    ws.merge_cells(f"A{cur_row}:{last_col}{cur_row}")
    c = ws[f"A{cur_row}"]
    c.value = subtitle
    c.font      = Font(name="Calibri", size=10, color="FFFFFF")
    c.fill      = _xl_fill(XL_SUBHEAD)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[cur_row].height = 20
    cur_row += 1

    ws.merge_cells(f"A{cur_row}:{last_col}{cur_row}")
    c = ws[f"A{cur_row}"]
    c.value = f"Generated by {REPORT_AUTHOR} on {datetime.now().strftime('%d %B %Y, %H:%M')}"
    c.font      = Font(name="Calibri", size=9, color="888888")
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[cur_row].height = 16
    cur_row += 1

    return cur_row


# =============================================================
# EXPORT ENGINE
# =============================================================

class ExportEngine:

    def __init__(self, station_config: dict = None):
        self.station_config = station_config or {}

        if station_config and "report_dirs" in station_config:
            self.pdf_dir  = station_config["report_dirs"]["pdf"]
            self.docx_dir = station_config["report_dirs"]["docx"]
            self.xlsx_dir = station_config["report_dirs"]["xlsx"]
        else:
            self.pdf_dir  = REPORTS_PDF_DIR
            self.docx_dir = REPORTS_DOCX_DIR
            self.xlsx_dir = REPORTS_PDF_DIR.parent / "xlsx"

        if station_config:
            self.station_name     = station_config.get("station_name", STATION_NAME)
            self.station_location = station_config.get("location", STATION_LOCATION)
            self.currency         = station_config.get("currency", REPORT_CURRENCY)
        else:
            self.station_name     = STATION_NAME
            self.station_location = STATION_LOCATION
            self.currency         = REPORT_CURRENCY

        self.pdf_dir.mkdir(parents=True, exist_ok=True)
        self.docx_dir.mkdir(parents=True, exist_ok=True)
        self.xlsx_dir.mkdir(parents=True, exist_ok=True)

    # =========================================================
    # MONTHLY PDF
    # =========================================================

    def generate_pdf(self, metrics, report_sections, period_label):
        period_label = _clean_period(period_label)
        filename = _filename("StationDeck_Report", period_label, "pdf")
        filepath = self.pdf_dir / filename

        doc = SimpleDocTemplate(
            str(filepath), pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )
        styles = self._pdf_styles()
        story  = []

        # Cover
        story += self._pdf_cover(styles, period_label, "OPERATIONAL PERFORMANCE REPORT")
        story.append(PageBreak())

        # Section 1 — Executive Summary
        story += self._pdf_section(styles, "1. Executive Summary",
                                   report_sections.get("executive_summary", ""))

        # Section 2 — Fuel Sales
        story += self._pdf_fuel_table(styles, metrics)
        story += self._pdf_section(styles, "2. Fuel Sales and Volume Analysis",
                                   report_sections.get("fuel_sales_analysis", ""))

        # Section 3 — Fuel Stock Movement
        if metrics.get("stock_data_available"):
            story += self._pdf_fuel_stock_table(styles, metrics)
            story += self._pdf_delivery_table(styles, metrics)
        story += self._pdf_section(styles, "3. Fuel Stock Movement and Deliveries",
                                   report_sections.get("fuel_stock_movement", ""))

        # Section 4 — Non-Fuel Products
        story += self._pdf_product_table(styles, metrics)
        story += self._pdf_section(styles, "4. Non-Fuel Product Performance",
                                   report_sections.get("product_performance", ""))

        # Section 5 — Shop
        story += self._pdf_section(styles, "5. Shop Performance",
                                   report_sections.get("shop_analysis", ""))

        # Section 6 — Payment Collection
        story += self._pdf_payment_table(styles, metrics)
        story += self._pdf_section(styles, "6. Payment Collection Analysis",
                                   report_sections.get("payment_collection_analysis", ""))

        # Section 7 — P&L
        story += self._pdf_expense_table(styles, metrics)
        story += self._pdf_section(styles, "7. Profit and Loss Analysis",
                                   report_sections.get("expense_pnl_analysis", ""))

        # Section 8 — Reconciliation
        story += self._pdf_reconciliation_table(styles, metrics)
        story += self._pdf_section(styles, "8. Cash Reconciliation and Operational Integrity",
                                   report_sections.get("cash_reconciliation_analysis", ""))

        # Sections 9–11
        story += self._pdf_section(styles, "9. Debtors and Depositors Analysis",
                                   report_sections.get("debtors_depositors_analysis", ""))
        story += self._pdf_section(styles, "10. Financial Position",
                                   report_sections.get("financial_position_analysis", ""))
        story += self._pdf_section(styles, "11. Claims and Outstanding Obligations",
                                   report_sections.get("claims_analysis", ""))

        # Footer
        story.append(Spacer(1, 0.5*cm))
        story.append(HRFlowable(width="100%", thickness=1, color=BRAND_MID))
        story.append(Paragraph(
            f"Generated by {REPORT_AUTHOR} \u00b7 "
            f"{datetime.now().strftime('%d %B %Y, %H:%M')}",
            styles["footer"]
        ))

        doc.build(story)
        print(f"  OK  PDF saved: {filepath}")
        return str(filepath)

    # =========================================================
    # ANNUAL PDF
    # =========================================================

    def generate_annual_pdf(self, metrics, report_sections, period_label):
        fy_label = metrics.get("fy_label", period_label)
        filename = _annual_filename("StationDeck_Report", fy_label, "pdf")
        filepath = self.pdf_dir / filename

        doc = SimpleDocTemplate(
            str(filepath), pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )
        styles = self._pdf_styles()
        story  = []

        story += self._pdf_cover(styles, fy_label, "ANNUAL OPERATIONAL PERFORMANCE REPORT",
                                 subtitle=period_label)
        story.append(PageBreak())

        section_map = [
            ("1. Annual Executive Summary",           "annual_executive_summary"),
            ("2. Annual Fuel Performance",            "annual_fuel_performance"),
            ("3. Annual Revenue Breakdown",           "annual_revenue_breakdown"),
            ("4. Annual Payment Trends",              "annual_payment_trends"),
            ("5. Annual Expense Analysis",            "annual_expense_analysis"),
            ("6. Annual Cash Reconciliation",         "annual_reconciliation"),
            ("7. Annual Outlook and Recommendations", "annual_outlook"),
        ]

        for title, key in section_map:
            story += self._pdf_section(styles, title, report_sections.get(key, ""))

        if metrics.get("monthly_breakdown"):
            story += self._pdf_monthly_breakdown_table(styles, metrics)

        if metrics.get("stock_data_available"):
            story += self._pdf_annual_fuel_table(styles, metrics)

        story.append(Spacer(1, 0.5*cm))
        story.append(HRFlowable(width="100%", thickness=1, color=BRAND_MID))
        story.append(Paragraph(
            f"Generated by {REPORT_AUTHOR} \u00b7 "
            f"{datetime.now().strftime('%d %B %Y, %H:%M')}",
            styles["footer"]
        ))

        doc.build(story)
        print(f"  OK  Annual PDF saved: {filepath}")
        return str(filepath)

    # =========================================================
    # PDF STYLES
    # =========================================================

    def _pdf_styles(self):
        s = {}
        s["cover_station"]    = ParagraphStyle("cover_station",   fontName="Helvetica-Bold", fontSize=26, textColor=white,               alignment=TA_CENTER, spaceAfter=6)
        s["cover_location"]   = ParagraphStyle("cover_location",  fontName="Helvetica",      fontSize=13, textColor=HexColor("#DDDDDD"),  alignment=TA_CENTER, spaceAfter=4)
        s["cover_title"]      = ParagraphStyle("cover_title",     fontName="Helvetica-Bold", fontSize=15, textColor=white,               alignment=TA_CENTER, spaceAfter=4)
        s["cover_subtitle"]   = ParagraphStyle("cover_subtitle",  fontName="Helvetica",      fontSize=11, textColor=HexColor("#BBBBBB"),  alignment=TA_CENTER, spaceAfter=4)
        s["cover_period"]     = ParagraphStyle("cover_period",    fontName="Helvetica-Bold", fontSize=22, textColor=white,               alignment=TA_CENTER, spaceAfter=8)
        s["cover_generated"]  = ParagraphStyle("cover_generated", fontName="Helvetica",      fontSize=10, textColor=HexColor("#AAAAAA"),  alignment=TA_CENTER)
        s["section_heading"]  = ParagraphStyle("section_heading", fontName="Helvetica-Bold", fontSize=12, textColor=BRAND_RED,           spaceBefore=14, spaceAfter=6)
        s["body"]             = ParagraphStyle("body",            fontName="Helvetica",      fontSize=10, textColor=TEXT_DARK,           leading=15, spaceAfter=8)
        s["table_header"]     = ParagraphStyle("table_header",    fontName="Helvetica-Bold", fontSize=9,  textColor=white,               alignment=TA_LEFT)
        s["table_cell"]       = ParagraphStyle("table_cell",      fontName="Helvetica",      fontSize=9,  textColor=TEXT_DARK)
        s["table_cell_right"] = ParagraphStyle("table_cell_right",fontName="Helvetica",      fontSize=9,  textColor=TEXT_DARK,           alignment=TA_RIGHT)
        s["footer"]           = ParagraphStyle("footer",          fontName="Helvetica",      fontSize=8,  textColor=HexColor("#888888"),  alignment=TA_CENTER, spaceBefore=4)
        s["sub_heading"]      = ParagraphStyle("sub_heading",     fontName="Helvetica-Bold", fontSize=10, textColor=BRAND_DARK,          spaceBefore=8, spaceAfter=4)
        return s

    def _pdf_cover(self, styles, period_label, report_type, subtitle=None):
        elements = []
        rows = [
            [Paragraph(self.station_name, styles["cover_station"])],
            [Paragraph(self.station_location, styles["cover_location"])],
            [Spacer(1, 0.8*cm)],
            [Paragraph(report_type, styles["cover_title"])],
        ]
        if subtitle:
            rows.append([Paragraph(subtitle, styles["cover_subtitle"])])
        rows += [
            [Paragraph(period_label.upper(), styles["cover_period"])],
            [Spacer(1, 0.5*cm)],
            [Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}", styles["cover_generated"])],
            [Paragraph(f"Prepared by: {REPORT_AUTHOR}", styles["cover_generated"])],
        ]
        ct = Table(rows, colWidths=[17*cm])
        ct.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,-1), BRAND_DARK),
            ("TOPPADDING",    (0,0), (-1,-1), 10),
            ("BOTTOMPADDING", (0,0), (-1,-1), 10),
            ("LEFTPADDING",   (0,0), (-1,-1), 20),
            ("RIGHTPADDING",  (0,0), (-1,-1), 20),
            ("ROUNDEDCORNERS",[8]),
        ]))
        elements.append(Spacer(1, 3*cm))
        elements.append(ct)
        elements.append(Spacer(1, 1*cm))
        elements.append(HRFlowable(width="100%", thickness=4, color=BRAND_RED))
        return elements

    def _pdf_section(self, styles, title, body):
        elements = []
        elements.append(Paragraph(title.upper(), styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        if body:
            for para in body.strip().split("\n\n"):
                if para.strip():
                    elements.append(Paragraph(para.strip(), styles["body"]))
        return elements

    def _pdf_fuel_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("FUEL PERFORMANCE SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        data = [
            [Paragraph("Metric", H),           Paragraph("PMS (Petrol)", H),                                                    Paragraph("AGO (Diesel)", H)],
            [Paragraph("Volume Sold", C),       Paragraph(_fmt(_safe(metrics,"pms_volume_total"),False)+" L", R),                Paragraph(_fmt(_safe(metrics,"ago_volume_total"),False)+" L", R)],
            [Paragraph("Revenue", C),           Paragraph(_fmt(_safe(metrics,"pms_revenue_total"),currency_symbol=cur), R),      Paragraph(_fmt(_safe(metrics,"ago_revenue_total"),currency_symbol=cur), R)],
            [Paragraph("Avg Daily Revenue", C), Paragraph(_fmt(_safe(metrics,"avg_daily_fuel_revenue"),currency_symbol=cur), R), Paragraph("", C)],
            [Paragraph("Combined Revenue", C),  Paragraph("", C),                                                                Paragraph(_fmt(_safe(metrics,"total_fuel_revenue"),currency_symbol=cur), R)],
        ]
        t = Table(data, colWidths=[6*cm, 5.5*cm, 5.5*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_fuel_stock_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("FUEL STOCK MOVEMENT", styles["sub_heading"]))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        fs  = metrics.get("fuel_stock", {})
        pms = fs.get("pms", {})
        ago = fs.get("ago", {})

        data = [
            [Paragraph("Item", H),                Paragraph("PMS (Petrol)", H),                                                                    Paragraph("AGO (Diesel)", H)],
            [Paragraph("Opening Dip", C),         Paragraph(_vol(pms.get("opening_dip_ltrs",0)), R),                                               Paragraph(_vol(ago.get("opening_dip_ltrs",0)), R)],
            [Paragraph("Total Purchases", C),     Paragraph(_vol(pms.get("total_purchases_ltrs",0)), R),                                           Paragraph(_vol(ago.get("total_purchases_ltrs",0)), R)],
            [Paragraph("Total Sales", C),         Paragraph(_vol(pms.get("total_sales_ltrs",0)), R),                                               Paragraph(_vol(ago.get("total_sales_ltrs",0)), R)],
            [Paragraph("Closing Dip", C),         Paragraph(_vol(pms.get("closing_dip_ltrs",0)), R),                                               Paragraph(_vol(ago.get("closing_dip_ltrs",0)), R)],
            [Paragraph("Loss / Gain", C),         Paragraph(_vol(pms.get("loss_gain_ltrs",0)), R),                                                 Paragraph(_vol(ago.get("loss_gain_ltrs",0)), R)],
            [Paragraph("Loss / Gain Value", C),   Paragraph(_fmt(pms.get("loss_gain_value_ugx",0),currency_symbol=cur), R),                        Paragraph(_fmt(ago.get("loss_gain_value_ugx",0),currency_symbol=cur), R)],
            [Paragraph("Avg Cost Price", C),      Paragraph(_fmt(pms.get("avg_cost_price_ugx",0),currency_symbol=cur)+"/L", R),                    Paragraph(_fmt(ago.get("avg_cost_price_ugx",0),currency_symbol=cur)+"/L", R)],
            [Paragraph("Avg Selling Price", C),   Paragraph(_fmt(pms.get("avg_selling_price_ugx",0),currency_symbol=cur)+"/L", R),                 Paragraph(_fmt(ago.get("avg_selling_price_ugx",0),currency_symbol=cur)+"/L", R)],
            [Paragraph("Closing Stock Value", C), Paragraph(_fmt(pms.get("closing_stock_value_ugx",0),currency_symbol=cur), R),                    Paragraph(_fmt(ago.get("closing_stock_value_ugx",0),currency_symbol=cur), R)],
        ]
        pms_loss  = pms.get("loss_gain_ltrs", 0)
        loss_color = SURPLUS_GREEN if pms_loss > 0 else DEFICIT_RED

        t = Table(data, colWidths=[6*cm, 5.5*cm, 5.5*cm])
        style = self._base_table_style()
        t.setStyle(TableStyle([
            *style._cmds,
            ("TEXTCOLOR", (1, 5), (2, 5), loss_color),
            ("FONTNAME",  (1, 5), (2, 5), "Helvetica-Bold"),
            ("TEXTCOLOR", (1, 6), (2, 6), loss_color),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 0.3*cm))
        return elements

    def _pdf_delivery_table(self, styles, metrics):
        fs  = metrics.get("fuel_stock", {})
        pms = fs.get("pms", {})
        ago = fs.get("ago", {})
        all_deliveries = []
        for d in pms.get("deliveries", []):
            all_deliveries.append(("PMS", d))
        for d in ago.get("deliveries", []):
            all_deliveries.append(("AGO", d))

        if not all_deliveries:
            return []

        all_deliveries.sort(key=lambda x: x[1]["date"])

        elements = []
        elements.append(Paragraph("FUEL DELIVERY LOG", styles["sub_heading"]))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency

        rows = [[Paragraph("Date", H), Paragraph("Product", H),
                 Paragraph("Litres", H), Paragraph("Cost/Litre", H), Paragraph("Total Value", H)]]
        for product, d in all_deliveries:
            rows.append([
                Paragraph(d["date"], C),
                Paragraph(product, C),
                Paragraph(f"{int(d['litres']):,} L", R),
                Paragraph(_fmt(d["cost_price"], currency_symbol=cur), R),
                Paragraph(_fmt(d["total_value"], currency_symbol=cur), R),
            ])

        t = Table(rows, colWidths=[3*cm, 2.5*cm, 3*cm, 4*cm, 4.5*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)
        elements.append(Spacer(1, 0.3*cm))
        return elements

    def _pdf_product_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("NON-FUEL PRODUCT REVENUE", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        ps  = metrics.get("product_sales", {})

        rows = [[Paragraph("Product", H), Paragraph("Revenue", H)]]
        product_items = [
            ("Lubricants",      ps.get("lubricants_ugx",    metrics.get("lubes_revenue_total", 0))),
            ("LPG Gas",         ps.get("lpg_ugx",           metrics.get("lpg_revenue_total", 0))),
            ("LPG Accessories", ps.get("lpg_accessories_ugx",metrics.get("lpg_accessories_total", 0))),
            ("TBA Credits",     ps.get("tba_ugx",           metrics.get("tba_revenue_total", 0))),
            ("Car Wash",        ps.get("car_wash_ugx",      metrics.get("car_wash_total", 0))),
            ("Shop Sales",      ps.get("shop_ugx",          metrics.get("shop_sales_total", 0))),
        ]
        for label, val in product_items:
            if float(val) > 0:
                rows.append([Paragraph(label, C), Paragraph(_fmt(val, currency_symbol=cur), R)])

        total_nonfuel = metrics.get("total_nonfuel_revenue", 0)
        rows.append([Paragraph("Total Non-Fuel Revenue", C),
                     Paragraph(_fmt(total_nonfuel, currency_symbol=cur), R)])

        t = Table(rows, colWidths=[9*cm, 8*cm])
        style = self._base_table_style()
        t.setStyle(TableStyle([
            *style._cmds,
            ("FONTNAME",  (0, len(rows)-1), (-1, len(rows)-1), "Helvetica-Bold"),
            ("BACKGROUND",(0, len(rows)-1), (-1, len(rows)-1), BRAND_LIGHT),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_payment_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("PAYMENT COLLECTION BREAKDOWN", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        data = [
            [Paragraph("Payment Type", H),         Paragraph("Amount", H),                                                          Paragraph("Share", H)],
            [Paragraph("Cash Collected", C),        Paragraph(_fmt(_safe(metrics,"cash_collected"),    currency_symbol=cur), R),     Paragraph(_pct(_safe(metrics,"cash_percentage")), R)],
            [Paragraph("Cashless Total", C),        Paragraph(_fmt(_safe(metrics,"cashless_collected"), currency_symbol=cur), R),    Paragraph(_pct(_safe(metrics,"cashless_percentage")), R)],
            [Paragraph("  \u21b3 Plus Card", C),    Paragraph(_fmt(_safe(metrics,"plus_card_total"),   currency_symbol=cur), R),     Paragraph("", C)],
            [Paragraph("  \u21b3 Visa", C),         Paragraph(_fmt(_safe(metrics,"visa_total"),        currency_symbol=cur), R),     Paragraph("", C)],
            [Paragraph("  \u21b3 Credit Sales", C), Paragraph(_fmt(_safe(metrics,"credit_sales_total"),currency_symbol=cur), R),    Paragraph("", C)],
            [Paragraph("Total Sales", C),           Paragraph(_fmt(_safe(metrics,"total_sales"),       currency_symbol=cur), R),     Paragraph("100%", R)],
        ]
        t = Table(data, colWidths=[6*cm, 6*cm, 5*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_expense_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("EXPENSE SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency

        data = [
            [Paragraph("Item", H),                   Paragraph("Amount", H)],
            [Paragraph("Total Expenses", C),          Paragraph(_fmt(_safe(metrics,"total_expenses"),currency_symbol=cur), R)],
            [Paragraph("Total Revenue", C),           Paragraph(_fmt(_safe(metrics,"total_revenue"),currency_symbol=cur), R)],
            [Paragraph("Avg Daily Fuel Revenue", C),  Paragraph(_fmt(_safe(metrics,"avg_daily_fuel_revenue"),currency_symbol=cur), R)],
        ]

        pnl = metrics.get("pnl", {})
        if pnl.get("gross_income", 0):
            data.append([Paragraph("Gross Income", C),    Paragraph(_fmt(pnl.get("gross_income",0),currency_symbol=cur), R)])
        if pnl.get("net_profit", 0):
            data.append([Paragraph("Net Profit", C),      Paragraph(_fmt(pnl.get("net_profit",0),currency_symbol=cur), R)])
        if pnl.get("reserve_balance", 0):
            data.append([Paragraph("Reserve Balance", C), Paragraph(_fmt(pnl.get("reserve_balance",0),currency_symbol=cur), R)])

        t = Table(data, colWidths=[9*cm, 8*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)

        exp_detail = metrics.get("expenses_detail_stock", {})
        if exp_detail and exp_detail.get("total_expenses", 0) > 0:
            elements.append(Spacer(1, 0.2*cm))
            elements.append(Paragraph("EXPENSE DETAIL BREAKDOWN", styles["sub_heading"]))
            detail_data = [[Paragraph("Expense Item", H), Paragraph("Amount", H)]]
            exp_labels = {
                "salaries": "Salaries", "electricity": "Electricity",
                "generator": "Generator", "meals": "Meals", "water": "Water",
                "security": "Security", "transport": "Transport",
                "airtime": "Airtime/Data", "nssf": "NSSF",
                "maintenance": "Maintenance", "stationary": "Stationery",
                "sanitation": "Sanitation", "vat_tax": "VAT/Tax",
                "sundries": "Sundries",
            }
            seen = set()
            for key, label in exp_labels.items():
                if label in seen:
                    continue
                val = exp_detail.get(key, 0)
                if float(val) > 0:
                    detail_data.append([Paragraph(label, C), Paragraph(_fmt(val,currency_symbol=cur), R)])
                    seen.add(label)
            if len(detail_data) > 1:
                dt = Table(detail_data, colWidths=[9*cm, 8*cm])
                dt.setStyle(self._base_table_style())
                elements.append(dt)

        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_reconciliation_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("CASH RECONCILIATION SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        delta_status = str(_safe(metrics, "delta_status", "UNKNOWN"))
        delta_color  = SURPLUS_GREEN if delta_status == "SURPLUS" else DEFICIT_RED
        data = [
            [Paragraph("Item", H),               Paragraph("Value", H)],
            [Paragraph("Cash Collected", C),      Paragraph(_fmt(_safe(metrics,"cash_collected"),currency_symbol=cur), R)],
            [Paragraph("Total Banked", C),        Paragraph(_fmt(_safe(metrics,"total_cash_banked"),currency_symbol=cur), R)],
            [Paragraph("Expected to Bank", C),    Paragraph(_fmt(_safe(metrics,"total_cash_expected"),currency_symbol=cur), R)],
            [Paragraph("Net Delta", C),           Paragraph(_fmt(_safe(metrics,"total_delta"),currency_symbol=cur), R)],
            [Paragraph("Delta Status", C),        Paragraph(delta_status, R)],
            [Paragraph("Days with Anomalies", C), Paragraph(str(_safe(metrics,"anomaly_days_count")), R)],
        ]
        t = Table(data, colWidths=[9*cm, 8*cm])
        style = self._base_table_style()
        t.setStyle(TableStyle([
            *style._cmds,
            ("TEXTCOLOR", (1,5),(1,5), delta_color),
            ("FONTNAME",  (1,5),(1,5), "Helvetica-Bold"),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_monthly_breakdown_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("MONTHLY PERFORMANCE BREAKDOWN", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency

        data = [[Paragraph("Month", H), Paragraph("Fuel Revenue", H),
                 Paragraph("Total Sales", H), Paragraph("Expenses", H), Paragraph("Delta", H)]]

        for row in metrics.get("monthly_breakdown", []):
            if row.get("data_available"):
                data.append([
                    Paragraph(row["label"], C),
                    Paragraph(_fmt(row.get("fuel_revenue",0), currency_symbol=cur), R),
                    Paragraph(_fmt(row.get("total_sales",0), currency_symbol=cur), R),
                    Paragraph(_fmt(row.get("expenses",0), currency_symbol=cur), R),
                    Paragraph(_fmt(row.get("delta",0), currency_symbol=cur), R),
                ])
            else:
                data.append([
                    Paragraph(row["label"], C),
                    Paragraph("—", R), Paragraph("—", R),
                    Paragraph("—", R), Paragraph("—", R),
                ])

        t = Table(data, colWidths=[3*cm, 4*cm, 4*cm, 3*cm, 3*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _pdf_annual_fuel_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("ANNUAL FUEL STOCK SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H, C, R = styles["table_header"], styles["table_cell"], styles["table_cell_right"]
        cur = self.currency
        fs  = metrics.get("fuel_stock", {})
        pms = fs.get("pms", {})
        ago = fs.get("ago", {})
        data = [
            [Paragraph("Metric", H),             Paragraph("PMS", H),                                                 Paragraph("AGO", H)],
            [Paragraph("Total Purchases", C),     Paragraph(_vol(pms.get("total_purchases_ltrs",0)), R),               Paragraph(_vol(ago.get("total_purchases_ltrs",0)), R)],
            [Paragraph("Total Sales", C),         Paragraph(_vol(pms.get("total_sales_ltrs",0)), R),                   Paragraph(_vol(ago.get("total_sales_ltrs",0)), R)],
            [Paragraph("Annual Loss/Gain", C),    Paragraph(_vol(pms.get("loss_gain_ltrs",0)), R),                     Paragraph(_vol(ago.get("loss_gain_ltrs",0)), R)],
            [Paragraph("Total Turnover", C),      Paragraph(_fmt(pms.get("turnover_ugx",0),currency_symbol=cur), R),   Paragraph(_fmt(ago.get("turnover_ugx",0),currency_symbol=cur), R)],
            [Paragraph("Delivery Count", C),      Paragraph(str(pms.get("delivery_count",0)), R),                      Paragraph(str(ago.get("delivery_count",0)), R)],
        ]
        t = Table(data, colWidths=[6*cm, 5.5*cm, 5.5*cm])
        t.setStyle(self._base_table_style())
        elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
        return elements

    def _base_table_style(self):
        return TableStyle([
            ("BACKGROUND",    (0,0), (-1, 0), BRAND_DARK),
            ("TEXTCOLOR",     (0,0), (-1, 0), white),
            ("FONTNAME",      (0,0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",      (0,0), (-1, 0), 9),
            ("TOPPADDING",    (0,0), (-1, 0), 7),
            ("BOTTOMPADDING", (0,0), (-1, 0), 7),
            ("FONTNAME",      (0,1), (-1,-1), "Helvetica"),
            ("FONTSIZE",      (0,1), (-1,-1), 9),
            ("TOPPADDING",    (0,1), (-1,-1), 5),
            ("BOTTOMPADDING", (0,1), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 8),
            ("RIGHTPADDING",  (0,0), (-1,-1), 8),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [white, BRAND_LIGHT]),
            ("GRID",          (0,0), (-1,-1), 0.5, BRAND_MID),
            ("LINEABOVE",     (0,0), (-1, 0), 1.5, BRAND_RED),
        ])

    # =========================================================
    # MONTHLY DOCX
    # =========================================================

    def generate_docx(self, metrics, report_sections, period_label):
        period_label = _clean_period(period_label)
        filename = _filename("StationDeck_Report", period_label, "docx")
        filepath = self.docx_dir / filename

        doc = Document()
        self._docx_set_margins(doc)
        self._docx_cover(doc, period_label, "OPERATIONAL PERFORMANCE REPORT")
        doc.add_page_break()

        sections = [
            ("Executive Summary",                      "executive_summary"),
            ("Fuel Sales and Volume Analysis",         "fuel_sales_analysis"),
            ("Fuel Stock Movement and Deliveries",     "fuel_stock_movement"),
            ("Non-Fuel Product Performance",           "product_performance"),
            ("Shop Performance",                       "shop_analysis"),
            ("Payment Collection Analysis",            "payment_collection_analysis"),
            ("Profit and Loss Analysis",               "expense_pnl_analysis"),
            ("Cash Reconciliation",                    "cash_reconciliation_analysis"),
            ("Debtors and Depositors Analysis",        "debtors_depositors_analysis"),
            ("Financial Position",                     "financial_position_analysis"),
            ("Claims and Outstanding Obligations",     "claims_analysis"),
        ]

        for i, (title, key) in enumerate(sections, 1):
            self._docx_section_heading(doc, f"{i}. {title}")

            if title == "Fuel Sales and Volume Analysis":
                self._docx_fuel_table(doc, metrics)
            elif title == "Fuel Stock Movement and Deliveries" and metrics.get("stock_data_available"):
                self._docx_fuel_stock_table(doc, metrics)
                self._docx_delivery_table(doc, metrics)
            elif title == "Non-Fuel Product Performance":
                self._docx_product_table(doc, metrics)
            elif title == "Payment Collection Analysis":
                self._docx_payment_table(doc, metrics)
            elif title == "Profit and Loss Analysis":
                self._docx_expense_table(doc, metrics)
            elif title == "Cash Reconciliation":
                self._docx_reconciliation_table(doc, metrics)

            self._docx_body(doc, report_sections.get(key, ""))

        self._docx_footer_line(doc, period_label)
        doc.save(str(filepath))
        print(f"  OK  DOCX saved: {filepath}")
        return str(filepath)

    # =========================================================
    # ANNUAL DOCX
    # =========================================================

    def generate_annual_docx(self, metrics, report_sections, period_label):
        fy_label = metrics.get("fy_label", period_label)
        filename = _annual_filename("StationDeck_Report", fy_label, "docx")
        filepath = self.docx_dir / filename

        doc = Document()
        self._docx_set_margins(doc)
        self._docx_cover(doc, fy_label, "ANNUAL OPERATIONAL PERFORMANCE REPORT",
                         subtitle=period_label)
        doc.add_page_break()

        annual_sections = [
            ("Annual Executive Summary",          "annual_executive_summary"),
            ("Annual Fuel Performance",           "annual_fuel_performance"),
            ("Annual Revenue Breakdown",          "annual_revenue_breakdown"),
            ("Annual Payment Trends",             "annual_payment_trends"),
            ("Annual Expense Analysis",           "annual_expense_analysis"),
            ("Annual Cash Reconciliation",        "annual_reconciliation"),
            ("Annual Outlook and Recommendations","annual_outlook"),
        ]

        for i, (title, key) in enumerate(annual_sections, 1):
            self._docx_section_heading(doc, f"{i}. {title}")
            if title == "Annual Fuel Performance" and metrics.get("stock_data_available"):
                self._docx_annual_fuel_table(doc, metrics)
            self._docx_body(doc, report_sections.get(key, ""))

        if metrics.get("monthly_breakdown"):
            self._docx_section_heading(doc, "Appendix: Monthly Performance Breakdown")
            self._docx_monthly_breakdown_table(doc, metrics)

        self._docx_footer_line(doc, fy_label)
        doc.save(str(filepath))
        print(f"  OK  Annual DOCX saved: {filepath}")
        return str(filepath)

    # =========================================================
    # DOCX HELPERS
    # =========================================================

    def _docx_set_margins(self, doc):
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2)
            section.right_margin  = Cm(2)

    def _docx_cover(self, doc, period_label, report_type, subtitle=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.station_name)
        run.font.size = Pt(28); run.font.bold = True; run.font.color.rgb = W_DARK

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.station_location)
        run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x88,0x88,0x88)

        doc.add_paragraph()

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_type)
        run.font.size = Pt(14); run.font.bold = True; run.font.color.rgb = W_RED

        if subtitle:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x88,0x88,0x88)

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(period_label.upper())
        run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = W_DARK

        doc.add_paragraph()

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
        run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x88,0x88,0x88)

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Prepared by: {REPORT_AUTHOR}")
        run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x88,0x88,0x88)

    def _docx_section_heading(self, doc, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(title.upper())
        run.font.size = Pt(12); run.font.bold = True; run.font.color.rgb = W_RED

    def _docx_sub_heading(self, doc, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(title)
        run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = W_DARK

    def _docx_body(self, doc, text):
        if not text:
            return
        for para in text.strip().split("\n\n"):
            if para.strip():
                p = doc.add_paragraph()
                run = p.add_run(para.strip())
                run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x22,0x22,0x22)
                p.paragraph_format.space_after = Pt(6)

    def _docx_fuel_table(self, doc, metrics):
        cur = self.currency
        self._docx_table(doc, ["Metric", "PMS (Petrol)", "AGO (Diesel)"], [
            ["Volume Sold",
             _fmt(_safe(metrics,"pms_volume_total"),False)+" L",
             _fmt(_safe(metrics,"ago_volume_total"),False)+" L"],
            ["Revenue",
             _fmt(_safe(metrics,"pms_revenue_total"),currency_symbol=cur),
             _fmt(_safe(metrics,"ago_revenue_total"),currency_symbol=cur)],
            ["Avg Daily Revenue",
             _fmt(_safe(metrics,"avg_daily_fuel_revenue"),currency_symbol=cur), ""],
            ["Combined Revenue", "",
             _fmt(_safe(metrics,"total_fuel_revenue"),currency_symbol=cur)],
        ])

    def _docx_fuel_stock_table(self, doc, metrics):
        cur = self.currency
        fs  = metrics.get("fuel_stock", {})
        pms = fs.get("pms", {})
        ago = fs.get("ago", {})
        self._docx_sub_heading(doc, "Fuel Stock Movement")
        self._docx_table(doc, ["Item", "PMS (Petrol)", "AGO (Diesel)"], [
            ["Opening Dip",         _vol(pms.get("opening_dip_ltrs",0)),           _vol(ago.get("opening_dip_ltrs",0))],
            ["Total Purchases",     _vol(pms.get("total_purchases_ltrs",0)),       _vol(ago.get("total_purchases_ltrs",0))],
            ["Total Sales",         _vol(pms.get("total_sales_ltrs",0)),           _vol(ago.get("total_sales_ltrs",0))],
            ["Closing Dip",         _vol(pms.get("closing_dip_ltrs",0)),           _vol(ago.get("closing_dip_ltrs",0))],
            ["Loss / Gain",         _vol(pms.get("loss_gain_ltrs",0)),             _vol(ago.get("loss_gain_ltrs",0))],
            ["Loss / Gain Value",   _fmt(pms.get("loss_gain_value_ugx",0),currency_symbol=cur), _fmt(ago.get("loss_gain_value_ugx",0),currency_symbol=cur)],
            ["Avg Cost Price",      _fmt(pms.get("avg_cost_price_ugx",0),currency_symbol=cur)+"/L", _fmt(ago.get("avg_cost_price_ugx",0),currency_symbol=cur)+"/L"],
            ["Avg Selling Price",   _fmt(pms.get("avg_selling_price_ugx",0),currency_symbol=cur)+"/L", _fmt(ago.get("avg_selling_price_ugx",0),currency_symbol=cur)+"/L"],
            ["Closing Stock Value", _fmt(pms.get("closing_stock_value_ugx",0),currency_symbol=cur), _fmt(ago.get("closing_stock_value_ugx",0),currency_symbol=cur)],
        ])

    def _docx_delivery_table(self, doc, metrics):
        cur = self.currency
        fs  = metrics.get("fuel_stock", {})
        all_deliveries = []
        for d in fs.get("pms", {}).get("deliveries", []):
            all_deliveries.append(("PMS", d))
        for d in fs.get("ago", {}).get("deliveries", []):
            all_deliveries.append(("AGO", d))
        if not all_deliveries:
            return
        all_deliveries.sort(key=lambda x: x[1]["date"])
        self._docx_sub_heading(doc, "Fuel Delivery Log")
        rows = []
        for product, d in all_deliveries:
            rows.append([
                d["date"], product,
                f"{int(d['litres']):,} L",
                _fmt(d["cost_price"],currency_symbol=cur),
                _fmt(d["total_value"],currency_symbol=cur),
            ])
        self._docx_table(doc, ["Date", "Product", "Litres", "Cost/Litre", "Total Value"], rows)

    def _docx_product_table(self, doc, metrics):
        cur = self.currency
        ps  = metrics.get("product_sales", {})
        rows = []
        items = [
            ("Lubricants",      ps.get("lubricants_ugx",    metrics.get("lubes_revenue_total",0))),
            ("LPG Gas",         ps.get("lpg_ugx",           metrics.get("lpg_revenue_total",0))),
            ("LPG Accessories", ps.get("lpg_accessories_ugx",metrics.get("lpg_accessories_total",0))),
            ("TBA Credits",     ps.get("tba_ugx",           metrics.get("tba_revenue_total",0))),
            ("Car Wash",        ps.get("car_wash_ugx",      metrics.get("car_wash_total",0))),
            ("Shop Sales",      ps.get("shop_ugx",          metrics.get("shop_sales_total",0))),
        ]
        for label, val in items:
            if float(val) > 0:
                rows.append([label, _fmt(val,currency_symbol=cur)])
        rows.append(["Total Non-Fuel Revenue",
                     _fmt(metrics.get("total_nonfuel_revenue",0),currency_symbol=cur)])
        self._docx_table(doc, ["Product", "Revenue"], rows)

    def _docx_payment_table(self, doc, metrics):
        cur = self.currency
        self._docx_table(doc, ["Payment Type", "Amount", "Share"], [
            ["Cash Collected",         _fmt(_safe(metrics,"cash_collected"),currency_symbol=cur),     _pct(_safe(metrics,"cash_percentage"))],
            ["Cashless Total",         _fmt(_safe(metrics,"cashless_collected"),currency_symbol=cur), _pct(_safe(metrics,"cashless_percentage"))],
            ["  \u21b3 Plus Card",    _fmt(_safe(metrics,"plus_card_total"),currency_symbol=cur),    ""],
            ["  \u21b3 Visa",         _fmt(_safe(metrics,"visa_total"),currency_symbol=cur),          ""],
            ["  \u21b3 Credit Sales", _fmt(_safe(metrics,"credit_sales_total"),currency_symbol=cur), ""],
            ["Total Sales",            _fmt(_safe(metrics,"total_sales"),currency_symbol=cur),        "100%"],
        ])

    def _docx_expense_table(self, doc, metrics):
        cur = self.currency
        pnl = metrics.get("pnl", {})
        rows = [
            ["Total Expenses",        _fmt(_safe(metrics,"total_expenses"),currency_symbol=cur)],
            ["Total Revenue",         _fmt(_safe(metrics,"total_revenue"),currency_symbol=cur)],
            ["Avg Daily Fuel Revenue",_fmt(_safe(metrics,"avg_daily_fuel_revenue"),currency_symbol=cur)],
        ]
        if pnl.get("gross_income",0):
            rows.append(["Gross Income",    _fmt(pnl["gross_income"],currency_symbol=cur)])
        if pnl.get("net_profit",0):
            rows.append(["Net Profit",      _fmt(pnl["net_profit"],currency_symbol=cur)])
        if pnl.get("reserve_balance",0):
            rows.append(["Reserve Balance", _fmt(pnl["reserve_balance"],currency_symbol=cur)])
        self._docx_table(doc, ["Item", "Amount"], rows)

    def _docx_reconciliation_table(self, doc, metrics):
        cur = self.currency
        self._docx_table(doc, ["Item", "Value"], [
            ["Cash Collected",      _fmt(_safe(metrics,"cash_collected"),currency_symbol=cur)],
            ["Total Banked",        _fmt(_safe(metrics,"total_cash_banked"),currency_symbol=cur)],
            ["Expected to Bank",    _fmt(_safe(metrics,"total_cash_expected"),currency_symbol=cur)],
            ["Net Delta",           _fmt(_safe(metrics,"total_delta"),currency_symbol=cur)],
            ["Delta Status",        str(_safe(metrics,"delta_status","UNKNOWN"))],
            ["Days with Anomalies", str(_safe(metrics,"anomaly_days_count"))],
        ])

    def _docx_annual_fuel_table(self, doc, metrics):
        cur = self.currency
        fs  = metrics.get("fuel_stock", {})
        pms = fs.get("pms", {})
        ago = fs.get("ago", {})
        self._docx_table(doc, ["Metric", "PMS", "AGO"], [
            ["Total Purchases",  _vol(pms.get("total_purchases_ltrs",0)),   _vol(ago.get("total_purchases_ltrs",0))],
            ["Total Sales",      _vol(pms.get("total_sales_ltrs",0)),       _vol(ago.get("total_sales_ltrs",0))],
            ["Annual Loss/Gain", _vol(pms.get("loss_gain_ltrs",0)),         _vol(ago.get("loss_gain_ltrs",0))],
            ["Total Turnover",   _fmt(pms.get("turnover_ugx",0),currency_symbol=cur), _fmt(ago.get("turnover_ugx",0),currency_symbol=cur)],
            ["Deliveries",       str(pms.get("delivery_count",0)),          str(ago.get("delivery_count",0))],
        ])

    def _docx_monthly_breakdown_table(self, doc, metrics):
        cur  = self.currency
        rows = []
        for row in metrics.get("monthly_breakdown", []):
            if row.get("data_available"):
                rows.append([
                    row["label"],
                    _fmt(row.get("fuel_revenue",0),currency_symbol=cur),
                    _fmt(row.get("total_sales",0),currency_symbol=cur),
                    _fmt(row.get("expenses",0),currency_symbol=cur),
                    _fmt(row.get("delta",0),currency_symbol=cur),
                ])
            else:
                rows.append([row["label"], "—", "—", "—", "—"])
        self._docx_table(doc, ["Month","Fuel Revenue","Total Sales","Expenses","Delta"], rows)

    def _docx_table(self, doc, headers, rows):
        table = doc.add_table(rows=1+len(rows), cols=len(headers))
        table.style = "Table Grid"
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = ""
            run = cell.paragraphs[0].add_run(headers[i])
            run.font.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            self._docx_set_cell_bg(cell, "1A1A2E")
        for r_idx, row_data in enumerate(rows):
            bg = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, cell in enumerate(table.rows[r_idx+1].cells):
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(row_data[c_idx]))
                run.font.size = Pt(9)
                self._docx_set_cell_bg(cell, bg)
        doc.add_paragraph()

    def _docx_set_cell_bg(self, cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _docx_footer_line(self, doc, period_label):
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"StationDeck Report \u00b7 {period_label} \u00b7 "
            f"Generated {datetime.now().strftime('%d %B %Y, %H:%M')} \u00b7 {REPORT_AUTHOR}"
        )
        run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x88,0x88,0x88)

    # =========================================================
    # MONTHLY XLSX
    # =========================================================

    def generate_xlsx(self, metrics, daily_df, period_label):
        period_label = _clean_period(period_label)
        filename     = _filename("StationDeck_Report", period_label, "xlsx")
        filepath     = self.xlsx_dir / filename

        wb = openpyxl.Workbook()
        self._xlsx_summary_sheet(wb, metrics, period_label)
        if metrics.get("stock_data_available"):
            self._xlsx_fuel_stock_sheet(wb, metrics, period_label)
        self._xlsx_product_sheet(wb, metrics, period_label)
        self._xlsx_daily_sheet(wb, daily_df, metrics, period_label)
        if "Sheet" in wb.sheetnames:
            del wb["Sheet"]
        wb.save(str(filepath))
        print(f"  OK  XLSX saved: {filepath}")
        return str(filepath)

    # =========================================================
    # ANNUAL XLSX
    # =========================================================

    def generate_annual_xlsx(self, metrics, period_label):
        fy_label = metrics.get("fy_label", period_label)
        filename = _annual_filename("StationDeck_Report", fy_label, "xlsx")
        filepath = self.xlsx_dir / filename

        wb = openpyxl.Workbook()
        self._xlsx_annual_summary_sheet(wb, metrics, fy_label)
        self._xlsx_annual_monthly_sheet(wb, metrics, fy_label)
        if metrics.get("stock_data_available"):
            self._xlsx_annual_fuel_sheet(wb, metrics, fy_label)
        if "Sheet" in wb.sheetnames:
            del wb["Sheet"]
        wb.save(str(filepath))
        print(f"  OK  Annual XLSX saved: {filepath}")
        return str(filepath)

    # =========================================================
    # XLSX SHEETS — MONTHLY
    # =========================================================

    def _xlsx_summary_sheet(self, wb, metrics, period_label):
        ws  = wb.create_sheet(title="Monthly Summary")
        cur = self.currency
        row = _xl_sheet_header(
            ws,
            f"{self.station_name} \u2014 {self.station_location}",
            f"Operational Performance Report \u2014 {period_label}",
            col_count=3,
        )

        def write_section(start_row, section_title, rows_data):
            ws.row_dimensions[start_row].height = 8
            start_row += 1
            ws.merge_cells(f"A{start_row}:C{start_row}")
            hdr = ws[f"A{start_row}"]
            hdr.value     = section_title
            hdr.font      = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
            hdr.fill      = _xl_fill(XL_RED)
            hdr.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            ws.row_dimensions[start_row].height = 18
            start_row += 1
            _xl_write_header_row(ws, start_row, ["Metric", "Value", ""], XL_DARK, height=16)
            start_row += 1
            for i, (label, value) in enumerate(rows_data):
                _xl_write_data_row(ws, start_row, [label, value, ""],
                                   alt=(i % 2 == 0), number_cols=[2])
                start_row += 1
            return start_row

        row = write_section(row, "FUEL PERFORMANCE", [
            ("PMS Volume Sold",     f"{_safe(metrics,'pms_volume_total',0):,.2f} L"),
            ("AGO Volume Sold",     f"{_safe(metrics,'ago_volume_total',0):,.2f} L"),
            ("PMS Revenue",         _fmt(_safe(metrics,"pms_revenue_total"),currency_symbol=cur)),
            ("AGO Revenue",         _fmt(_safe(metrics,"ago_revenue_total"),currency_symbol=cur)),
            ("Total Fuel Revenue",  _fmt(_safe(metrics,"total_fuel_revenue"),currency_symbol=cur)),
            ("Avg Daily Fuel Rev.", _fmt(_safe(metrics,"avg_daily_fuel_revenue"),currency_symbol=cur)),
        ])
        row = write_section(row, "NON-FUEL REVENUE", [
            ("Lubricants Revenue",   _fmt(_safe(metrics,"lubes_revenue_total"),currency_symbol=cur)),
            ("LPG Gas Revenue",      _fmt(_safe(metrics,"lpg_revenue_total"),currency_symbol=cur)),
            ("LPG Accessories",      _fmt(metrics.get("lpg_accessories_total",0),currency_symbol=cur)),
            ("TBA Credits",          _fmt(metrics.get("tba_revenue_total",0),currency_symbol=cur)),
            ("Car Wash",             _fmt(metrics.get("car_wash_total",0),currency_symbol=cur)),
            ("Shop Sales",           _fmt(_safe(metrics,"shop_sales_total"),currency_symbol=cur)),
            ("Total Non-Fuel Rev.",  _fmt(_safe(metrics,"total_nonfuel_revenue"),currency_symbol=cur)),
            ("Total Sales",          _fmt(_safe(metrics,"total_sales"),currency_symbol=cur)),
            ("Total Revenue",        _fmt(_safe(metrics,"total_revenue"),currency_symbol=cur)),
        ])
        row = write_section(row, "PAYMENT COLLECTION", [
            ("Cash Collected",       _fmt(_safe(metrics,"cash_collected"),currency_symbol=cur)),
            ("Cash Share",           _pct(_safe(metrics,"cash_percentage"))),
            ("Cashless Collected",   _fmt(_safe(metrics,"cashless_collected"),currency_symbol=cur)),
            ("Cashless Share",       _pct(_safe(metrics,"cashless_percentage"))),
            ("Plus Card Total",      _fmt(_safe(metrics,"plus_card_total"),currency_symbol=cur)),
            ("Visa Total",           _fmt(_safe(metrics,"visa_total"),currency_symbol=cur)),
            ("Credit Sales",         _fmt(_safe(metrics,"credit_sales_total"),currency_symbol=cur)),
        ])

        pnl = metrics.get("pnl", {})
        row = write_section(row, "PROFIT & LOSS", [
            ("Total Expenses",       _fmt(_safe(metrics,"total_expenses"),currency_symbol=cur)),
            ("Gross Income",         _fmt(pnl.get("gross_income",0),currency_symbol=cur)),
            ("Net Profit",           _fmt(pnl.get("net_profit",0),currency_symbol=cur)),
            ("Reserve Balance",      _fmt(pnl.get("reserve_balance",0),currency_symbol=cur)),
        ])

        delta_status = str(_safe(metrics,"delta_status","UNKNOWN"))
        delta_color  = XL_GREEN if delta_status == "SURPLUS" else XL_RED
        row = write_section(row, "CASH RECONCILIATION", [
            ("Days Covered",         str(_safe(metrics,"total_days"))),
            ("Cash Collected",       _fmt(_safe(metrics,"cash_collected"),currency_symbol=cur)),
            ("Total Banked",         _fmt(_safe(metrics,"total_cash_banked"),currency_symbol=cur)),
            ("Net Delta",            _fmt(_safe(metrics,"total_delta"),currency_symbol=cur)),
            ("Delta Status",         delta_status),
            ("Days with Anomalies",  str(_safe(metrics,"anomaly_days_count"))),
        ])
        ws[f"B{row-2}"].font = Font(name="Calibri", bold=True, size=10, color=delta_color)

        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 22
        ws.column_dimensions["C"].width = 4

    def _xlsx_fuel_stock_sheet(self, wb, metrics, period_label):
        ws  = wb.create_sheet(title="Fuel Stock Movement")
        cur = self.currency
        row = _xl_sheet_header(
            ws,
            f"{self.station_name} \u2014 Fuel Stock Movement",
            period_label, col_count=6,
        )

        fs  = metrics.get("fuel_stock", {})
        pms = fs.get("pms", {})
        ago = fs.get("ago", {})

        row += 1
        _xl_write_header_row(ws, row, ["Item", "PMS (Petrol)", "", "AGO (Diesel)", "", ""], XL_DARK)
        row += 1
        summary_rows = [
            ("Opening Dip (L)",     f"{pms.get('opening_dip_ltrs',0):,.2f}",      "", f"{ago.get('opening_dip_ltrs',0):,.2f}",     ""),
            ("Total Purchases (L)", f"{pms.get('total_purchases_ltrs',0):,.2f}",  "", f"{ago.get('total_purchases_ltrs',0):,.2f}", ""),
            ("Total Sales (L)",     f"{pms.get('total_sales_ltrs',0):,.2f}",      "", f"{ago.get('total_sales_ltrs',0):,.2f}",     ""),
            ("Closing Dip (L)",     f"{pms.get('closing_dip_ltrs',0):,.2f}",      "", f"{ago.get('closing_dip_ltrs',0):,.2f}",     ""),
            ("Loss / Gain (L)",     f"{pms.get('loss_gain_ltrs',0):,.2f}",        "", f"{ago.get('loss_gain_ltrs',0):,.2f}",       ""),
            ("Loss / Gain (UGX)",   _fmt(pms.get("loss_gain_value_ugx",0),currency_symbol=cur), "", _fmt(ago.get("loss_gain_value_ugx",0),currency_symbol=cur), ""),
            ("Avg Cost Price",      _fmt(pms.get("avg_cost_price_ugx",0),currency_symbol=cur)+"/L", "", _fmt(ago.get("avg_cost_price_ugx",0),currency_symbol=cur)+"/L", ""),
            ("Avg Selling Price",   _fmt(pms.get("avg_selling_price_ugx",0),currency_symbol=cur)+"/L", "", _fmt(ago.get("avg_selling_price_ugx",0),currency_symbol=cur)+"/L", ""),
            ("Closing Stock Value", _fmt(pms.get("closing_stock_value_ugx",0),currency_symbol=cur), "", _fmt(ago.get("closing_stock_value_ugx",0),currency_symbol=cur), ""),
            ("Delivery Count",      str(pms.get("delivery_count",0)),              "", str(ago.get("delivery_count",0)),             ""),
        ]
        for i, vals in enumerate(summary_rows):
            _xl_write_data_row(ws, row, list(vals), alt=(i%2==0))
            row += 1

        # Delivery log
        row += 2
        ws.merge_cells(f"A{row}:F{row}")
        hdr = ws[f"A{row}"]
        hdr.value = "DELIVERY LOG"
        hdr.font  = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
        hdr.fill  = _xl_fill(XL_ORANGE)
        hdr.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.row_dimensions[row].height = 18
        row += 1

        _xl_write_header_row(ws, row, ["Date", "Product", "Litres", "Cost/Litre", "Total Value", ""], XL_DARK)
        row += 1

        all_deliveries = []
        for d in pms.get("deliveries", []):
            all_deliveries.append(("PMS", d))
        for d in ago.get("deliveries", []):
            all_deliveries.append(("AGO", d))
        all_deliveries.sort(key=lambda x: x[1]["date"])

        for i, (product, d) in enumerate(all_deliveries):
            _xl_write_data_row(ws, row, [
                d["date"], product,
                f"{int(d['litres']):,} L",
                _fmt(d["cost_price"],currency_symbol=cur),
                _fmt(d["total_value"],currency_symbol=cur), "",
            ], alt=(i%2==0))
            row += 1

        for i, w in enumerate([16, 10, 14, 18, 18, 4], start=1):
            ws.column_dimensions[get_column_letter(i)].width = w
        ws.freeze_panes = "A4"

    def _xlsx_product_sheet(self, wb, metrics, period_label):
        ws  = wb.create_sheet(title="Product Sales")
        cur = self.currency
        row = _xl_sheet_header(
            ws,
            f"{self.station_name} \u2014 Non-Fuel Product Sales",
            period_label, col_count=3,
        )

        ps = metrics.get("product_sales", {})
        row += 1
        _xl_write_header_row(ws, row, ["Product", "Revenue (UGX)", ""], XL_DARK)
        row += 1

        items = [
            ("Lubricants",      ps.get("lubricants_ugx",    metrics.get("lubes_revenue_total",0))),
            ("LPG Gas",         ps.get("lpg_ugx",           metrics.get("lpg_revenue_total",0))),
            ("LPG Accessories", ps.get("lpg_accessories_ugx",metrics.get("lpg_accessories_total",0))),
            ("TBA Credits",     ps.get("tba_ugx",           metrics.get("tba_revenue_total",0))),
            ("Car Wash",        ps.get("car_wash_ugx",      metrics.get("car_wash_total",0))),
            ("Shop Sales",      ps.get("shop_ugx",          metrics.get("shop_sales_total",0))),
        ]
        for i, (label, val) in enumerate(items):
            _xl_write_data_row(ws, row, [label, _fmt(val,currency_symbol=cur), ""], alt=(i%2==0))
            row += 1

        # Total row
        ws.row_dimensions[row].height = 18
        for col in range(1, 4):
            cell = ws.cell(row=row, column=col)
            cell.fill   = _xl_fill(XL_DARK)
            cell.font   = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
            cell.border = _xl_border()
        ws.cell(row=row, column=1, value="Total Non-Fuel Revenue")
        ws.cell(row=row, column=2, value=_fmt(metrics.get("total_nonfuel_revenue",0),currency_symbol=cur))
        ws.cell(row=row, column=2).alignment = Alignment(horizontal="right", vertical="center")

        ws.column_dimensions["A"].width = 22
        ws.column_dimensions["B"].width = 22
        ws.column_dimensions["C"].width = 4

    def _xlsx_daily_sheet(self, wb, daily_df, metrics, period_label):
        """
        Daily Records sheet — combines cashflow daily data with
        non-fuel product columns from General Sales Summary.

        Columns:
          Date | PMS Vol | AGO Vol | PMS Rev | AGO Rev |
          Lubricants | LPG | LPG Acc | TBA | Car Wash | Shop |
          Total Sales | Cash | Cashless | Expenses | Delta
        """
        if daily_df is None or daily_df.empty:
            return

        ws = wb.create_sheet(title="Daily Records")

        # Column headers — 16 columns total
        headers = [
            "Date",
            "PMS Vol (L)", "AGO Vol (L)",
            "PMS Rev (UGX)", "AGO Rev (UGX)",
            "Lubricants (UGX)", "LPG Gas (UGX)", "LPG Acc (UGX)",
            "TBA (UGX)", "Car Wash (UGX)", "Shop Sales (UGX)",
            "Total Sales (UGX)", "Cash (UGX)", "Cashless (UGX)",
            "Expenses (UGX)", "Delta (UGX)",
        ]
        col_count = len(headers)  # 16

        row = _xl_sheet_header(
            ws,
            f"{self.station_name} \u2014 Daily Records \u2014 {period_label}",
            "", col_count=col_count,
        )

        _xl_write_header_row(ws, row, headers, XL_DARK, height=20)
        row += 1

        # Build a date-keyed lookup from general_sales_daily
        # Each entry: {"date": "01 Apr 2026", "lubricants_ugx": ..., ...}
        general_sales_lookup = {}
        for entry in metrics.get("general_sales_daily", []):
            general_sales_lookup[entry["date"]] = entry

        number_cols = list(range(2, col_count + 1))

        for i, (_, row_data) in enumerate(daily_df.iterrows()):
            alt = (i % 2 == 0)

            def safe_val(col_name, decimals=0, _r=row_data):
                try:
                    val = _r.get(col_name, 0)
                    return round(float(val), decimals) if decimals else int(float(val))
                except Exception:
                    return 0

            # Date string — used for lookup and display
            try:
                dv = row_data["date"]
                date_str = dv.strftime("%d %b %Y") if hasattr(dv, "strftime") else str(dv)
            except Exception:
                date_str = "N/A"

            # Get product sales for this day (0 if not available)
            gs = general_sales_lookup.get(date_str, {})

            cashless_val = safe_val("cashless_total") - safe_val("total_cash")

            values = [
                date_str,
                # Fuel volumes
                safe_val("pms_volume", 2),
                safe_val("ago_volume", 2),
                # Fuel revenues
                safe_val("pms_revenue"),
                safe_val("ago_revenue"),
                # Non-fuel products from General Sales Summary
                int(gs.get("lubricants_ugx", 0)),
                int(gs.get("lpg_ugx", 0)),
                int(gs.get("lpg_accessories_ugx", 0)),
                int(gs.get("tba_ugx", 0)),
                int(gs.get("car_wash_ugx", 0)),
                int(gs.get("shop_ugx", 0)),
                # Summary columns
                safe_val("cashless_total"),
                safe_val("total_cash"),
                cashless_val,
                safe_val("total_expenses"),
                safe_val("delta"),
            ]

            _xl_write_data_row(ws, row, values, alt=alt, number_cols=number_cols)

            # Colour delta column green/red
            delta_v = values[15]
            if isinstance(delta_v, (int, float)):
                ws.cell(row=row, column=16).font = Font(
                    name="Calibri", size=10,
                    color=XL_GREEN if delta_v >= 0 else XL_RED
                )

            row += 1

        # Column widths
        col_widths = [14, 13, 13, 16, 16, 16, 14, 14, 12, 13, 15, 18, 16, 16, 16, 14]
        for i, w in enumerate(col_widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w

        ws.freeze_panes = "A4"

    # =========================================================
    # XLSX SHEETS — ANNUAL
    # =========================================================

    def _xlsx_annual_summary_sheet(self, wb, metrics, fy_label):
        ws  = wb.create_sheet(title="Annual Summary")
        cur = self.currency
        row = _xl_sheet_header(
            ws,
            f"{self.station_name} \u2014 Annual Performance",
            f"{fy_label} (July {metrics.get('fy_start_year','')} – June {metrics.get('fy_start_year',0)+1})",
            col_count=3,
        )

        def write_section(start_row, title, rows_data):
            ws.row_dimensions[start_row].height = 8
            start_row += 1
            ws.merge_cells(f"A{start_row}:C{start_row}")
            hdr = ws[f"A{start_row}"]
            hdr.value = title
            hdr.font  = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
            hdr.fill  = _xl_fill(XL_RED)
            hdr.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            ws.row_dimensions[start_row].height = 18
            start_row += 1
            _xl_write_header_row(ws, start_row, ["Metric", "Value", ""], XL_DARK, height=16)
            start_row += 1
            for i, (label, value) in enumerate(rows_data):
                _xl_write_data_row(ws, start_row, [label, value, ""], alt=(i%2==0))
                start_row += 1
            return start_row

        row = write_section(row, "ANNUAL FUEL PERFORMANCE", [
            ("PMS Volume Sold",      f"{metrics.get('pms_volume_total',0):,.2f} L"),
            ("AGO Volume Sold",      f"{metrics.get('ago_volume_total',0):,.2f} L"),
            ("Total Fuel Revenue",   _fmt(metrics.get("total_fuel_revenue",0),currency_symbol=cur)),
            ("Avg Monthly Fuel Rev.",_fmt(metrics.get("avg_monthly_fuel_revenue",0),currency_symbol=cur)),
        ])
        row = write_section(row, "ANNUAL REVENUE", [
            ("Lubricants Revenue",   _fmt(metrics.get("lubes_revenue_total",0),currency_symbol=cur)),
            ("LPG Gas Revenue",      _fmt(metrics.get("lpg_revenue_total",0),currency_symbol=cur)),
            ("Shop Sales",           _fmt(metrics.get("shop_sales_total",0),currency_symbol=cur)),
            ("Total Revenue",        _fmt(metrics.get("total_revenue",0),currency_symbol=cur)),
        ])
        row = write_section(row, "ANNUAL PAYMENT TRENDS", [
            ("Cash Share",           _pct(metrics.get("cash_percentage",0))),
            ("Cashless Share",       _pct(metrics.get("cashless_percentage",0))),
            ("Total Expenses",       _fmt(metrics.get("total_expenses",0),currency_symbol=cur)),
        ])
        delta_status = str(metrics.get("delta_status","UNKNOWN"))
        delta_color  = XL_GREEN if delta_status == "SURPLUS" else XL_RED
        row = write_section(row, "ANNUAL RECONCILIATION", [
            ("Months with Data",     str(metrics.get("months_with_data",0))),
            ("Total Days",           str(metrics.get("total_days",0))),
            ("Annual Net Delta",     _fmt(metrics.get("total_delta",0),currency_symbol=cur)),
            ("Delta Status",         delta_status),
        ])
        ws[f"B{row-2}"].font = Font(name="Calibri", bold=True, size=10, color=delta_color)

        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 22
        ws.column_dimensions["C"].width = 4

    def _xlsx_annual_monthly_sheet(self, wb, metrics, fy_label):
        ws  = wb.create_sheet(title="Monthly Breakdown")
        cur = self.currency
        col_count = 7
        row = _xl_sheet_header(
            ws,
            f"{self.station_name} \u2014 Monthly Breakdown",
            fy_label, col_count=col_count,
        )
        row += 1
        _xl_write_header_row(ws, row, [
            "Month", "Days", "Fuel Revenue", "Total Sales",
            "Cash", "Expenses", "Delta"
        ], XL_DARK, height=20)
        row += 1

        for i, m in enumerate(metrics.get("monthly_breakdown", [])):
            alt = (i % 2 == 0)
            if m.get("data_available"):
                values = [
                    m["label"],
                    m.get("days", 0),
                    _fmt(m.get("fuel_revenue",0),currency_symbol=cur),
                    _fmt(m.get("total_sales",0),currency_symbol=cur),
                    _fmt(m.get("cash",0),currency_symbol=cur),
                    _fmt(m.get("expenses",0),currency_symbol=cur),
                    _fmt(m.get("delta",0),currency_symbol=cur),
                ]
                _xl_write_data_row(ws, row, values, alt=alt)
                delta_v = m.get("delta", 0)
                ws.cell(row=row, column=7).font = Font(
                    name="Calibri", size=10,
                    color=XL_GREEN if delta_v >= 0 else XL_RED
                )
            else:
                _xl_write_data_row(ws, row, [m["label"],"—","—","—","—","—","—"], alt=alt)
            row += 1

        for i, w in enumerate([14, 8, 20, 20, 18, 18, 18], start=1):
            ws.column_dimensions[get_column_letter(i)].width = w
        ws.freeze_panes = "A4"

    def _xlsx_annual_fuel_sheet(self, wb, metrics, fy_label):
        ws  = wb.create_sheet(title="Annual Fuel Stock")
        cur = self.currency
        row = _xl_sheet_header(
            ws,
            f"{self.station_name} \u2014 Annual Fuel Stock Summary",
            fy_label, col_count=3,
        )
        row += 1
        _xl_write_header_row(ws, row, ["Metric", "PMS (Petrol)", "AGO (Diesel)"], XL_DARK)
        row += 1

        fs  = metrics.get("fuel_stock", {})
        pms = fs.get("pms", {})
        ago = fs.get("ago", {})

        summary_rows = [
            ("Total Purchases (L)",  f"{pms.get('total_purchases_ltrs',0):,.2f}", f"{ago.get('total_purchases_ltrs',0):,.2f}"),
            ("Total Sales (L)",      f"{pms.get('total_sales_ltrs',0):,.2f}",     f"{ago.get('total_sales_ltrs',0):,.2f}"),
            ("Annual Loss/Gain (L)", f"{pms.get('loss_gain_ltrs',0):,.2f}",       f"{ago.get('loss_gain_ltrs',0):,.2f}"),
            ("Total Turnover",       _fmt(pms.get("turnover_ugx",0),currency_symbol=cur), _fmt(ago.get("turnover_ugx",0),currency_symbol=cur)),
            ("Total Purchase Value", _fmt(pms.get("purchase_value_ugx",0),currency_symbol=cur), _fmt(ago.get("purchase_value_ugx",0),currency_symbol=cur)),
            ("Delivery Count",       str(pms.get("delivery_count",0)),             str(ago.get("delivery_count",0))),
        ]
        for i, vals in enumerate(summary_rows):
            _xl_write_data_row(ws, row, list(vals), alt=(i%2==0))
            row += 1

        ws.column_dimensions["A"].width = 24
        ws.column_dimensions["B"].width = 22
        ws.column_dimensions["C"].width = 22