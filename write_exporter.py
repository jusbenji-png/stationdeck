# Run this from your stationdeck root folder:
#   python write_exporter.py
# It will overwrite src/exporter.py with the complete correct version.

from pathlib import Path

dest = Path(__file__).parent / "src" / "exporter.py"

content = r'''# =============================================================
# StationDeck - Export Engine
# =============================================================
from pathlib import Path
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor, white, black
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from config.settings import (
    REPORTS_PDF_DIR,
    REPORTS_DOCX_DIR,
    STATION_NAME,
    STATION_LOCATION,
    REPORT_CURRENCY,
    REPORT_AUTHOR,
)

REPORTS_XLSX_DIR = REPORTS_PDF_DIR.parent / "xlsx"

BRAND_RED     = HexColor("#C8102E")
BRAND_DARK    = HexColor("#1A1A2E")
BRAND_LIGHT   = HexColor("#F5F5F5")
BRAND_MID     = HexColor("#E0E0E0")
TEXT_DARK     = HexColor("#222222")
SURPLUS_GREEN = HexColor("#1E7E34")
DEFICIT_RED   = HexColor("#C8102E")

W_RED  = RGBColor(0xC8, 0x10, 0x2E)
W_DARK = RGBColor(0x1A, 0x1A, 0x2E)
W_GREY = RGBColor(0xF5, 0xF5, 0xF5)
W_MID  = RGBColor(0xE0, 0xE0, 0xE0)

XL_DARK    = "1A1A2E"
XL_RED     = "C8102E"
XL_LIGHT   = "F5F5F5"
XL_WHITE   = "FFFFFF"
XL_MID     = "E0E0E0"
XL_GREEN   = "1E7E34"
XL_SUBHEAD = "2E4057"


def _fmt(value, currency=True):
    try:
        formatted = f"{int(value):,}"
        if currency:
            return f"{REPORT_CURRENCY} {formatted}"
        return formatted
    except (TypeError, ValueError):
        return str(value)


def _pct(value):
    try:
        return f"{float(value):.1f}%"
    except (TypeError, ValueError):
        return str(value)


def _safe(metrics, key, default="N/A"):
    return metrics.get(key, default)


def _filename(prefix, period_label, extension):
    clean = period_label.replace(" ", "_")
    return f"{prefix}_{clean}.{extension}"


def _xl_fill(hex_color):
    return PatternFill(fill_type="solid", fgColor=hex_color)


def _xl_border(color="CCCCCC"):
    side = Side(style="thin", color=color)
    return Border(left=side, right=side, top=side, bottom=side)


def _xl_header_font(color="FFFFFF", size=10, bold=True):
    return Font(name="Calibri", bold=bold, size=size, color=color)


def _xl_body_font(size=10, bold=False, color="222222"):
    return Font(name="Calibri", size=size, bold=bold, color=color)


def _xl_write_header_row(ws, row, values, fill_color, font_color="FFFFFF", height=20):
    ws.row_dimensions[row].height = height
    for col, value in enumerate(values, start=1):
        cell = ws.cell(row=row, column=col, value=value)
        cell.fill      = _xl_fill(fill_color)
        cell.font      = _xl_header_font(color=font_color)
        cell.border    = _xl_border()
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)


def _xl_write_data_row(ws, row, values, alt=False, bold=False, number_cols=None):
    fill = _xl_fill(XL_LIGHT) if alt else _xl_fill(XL_WHITE)
    number_cols = number_cols or []
    ws.row_dimensions[row].height = 18
    for col, value in enumerate(values, start=1):
        cell = ws.cell(row=row, column=col, value=value)
        cell.fill      = fill
        cell.font      = _xl_body_font(bold=bold)
        cell.border    = _xl_border()
        align          = "right" if col in number_cols else "left"
        cell.alignment = Alignment(horizontal=align, vertical="center")


class ExportEngine:

    def __init__(self):
        REPORTS_PDF_DIR.mkdir(parents=True, exist_ok=True)
        REPORTS_DOCX_DIR.mkdir(parents=True, exist_ok=True)
        REPORTS_XLSX_DIR.mkdir(parents=True, exist_ok=True)

    # =========================================================
    # PDF
    # =========================================================

    def generate_pdf(self, metrics, report_sections, period_label):
        filename = _filename("StationDeck_Report", period_label, "pdf")
        filepath = REPORTS_PDF_DIR / filename
        doc = SimpleDocTemplate(str(filepath), pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        styles = self._pdf_styles()
        story  = []
        story += self._pdf_cover(styles, period_label)
        story.append(PageBreak())
        story += self._pdf_section(styles, "Executive Summary",       report_sections.get("executive_summary", ""))
        story += self._pdf_fuel_table(styles, metrics)
        story += self._pdf_section(styles, "Fuel Sales Analysis",     report_sections.get("fuel_sales_analysis", ""))
        story += self._pdf_payment_table(styles, metrics)
        story += self._pdf_section(styles, "Payment Collection Analysis", report_sections.get("payment_collection_analysis", ""))
        story += self._pdf_expense_table(styles, metrics)
        story += self._pdf_reconciliation_table(styles, metrics)
        story += self._pdf_section(styles, "Cash Reconciliation Analysis", report_sections.get("cash_reconciliation_analysis", ""))
        story.append(Spacer(1, 0.5*cm))
        story.append(HRFlowable(width="100%", thickness=1, color=BRAND_MID))
        story.append(Paragraph(
            f"Generated by {REPORT_AUTHOR} \u00b7 {datetime.now().strftime('%d %B %Y, %H:%M')}",
            styles["footer"]))
        doc.build(story)
        print(f"  \u2705 PDF saved: {filepath}")
        return str(filepath)

    def _pdf_styles(self):
        styles = {}
        styles["cover_station"]   = ParagraphStyle("cover_station",  fontName="Helvetica-Bold", fontSize=26, textColor=white, alignment=TA_CENTER, spaceAfter=6)
        styles["cover_location"]  = ParagraphStyle("cover_location", fontName="Helvetica", fontSize=13, textColor=HexColor("#DDDDDD"), alignment=TA_CENTER, spaceAfter=4)
        styles["cover_title"]     = ParagraphStyle("cover_title",    fontName="Helvetica-Bold", fontSize=16, textColor=white, alignment=TA_CENTER, spaceAfter=4)
        styles["cover_period"]    = ParagraphStyle("cover_period",   fontName="Helvetica-Bold", fontSize=22, textColor=white, alignment=TA_CENTER, spaceAfter=8)
        styles["cover_generated"] = ParagraphStyle("cover_generated",fontName="Helvetica", fontSize=10, textColor=HexColor("#AAAAAA"), alignment=TA_CENTER)
        styles["section_heading"] = ParagraphStyle("section_heading",fontName="Helvetica-Bold", fontSize=13, textColor=BRAND_RED, spaceBefore=14, spaceAfter=6)
        styles["body"]            = ParagraphStyle("body",           fontName="Helvetica", fontSize=10, textColor=TEXT_DARK, leading=15, spaceAfter=8)
        styles["table_header"]    = ParagraphStyle("table_header",   fontName="Helvetica-Bold", fontSize=9, textColor=white, alignment=TA_LEFT)
        styles["table_cell"]      = ParagraphStyle("table_cell",     fontName="Helvetica", fontSize=9, textColor=TEXT_DARK)
        styles["table_cell_right"]= ParagraphStyle("table_cell_right",fontName="Helvetica", fontSize=9, textColor=TEXT_DARK, alignment=TA_RIGHT)
        styles["footer"]          = ParagraphStyle("footer",         fontName="Helvetica", fontSize=8, textColor=HexColor("#888888"), alignment=TA_CENTER, spaceBefore=4)
        return styles

    def _pdf_cover(self, styles, period_label):
        elements = []
        cover_data = [
            [Paragraph(STATION_NAME, styles["cover_station"])],
            [Paragraph(STATION_LOCATION, styles["cover_location"])],
            [Spacer(1, 1*cm)],
            [Paragraph("OPERATIONAL PERFORMANCE REPORT", styles["cover_title"])],
            [Paragraph(period_label.upper(), styles["cover_period"])],
            [Spacer(1, 0.5*cm)],
            [Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}", styles["cover_generated"])],
            [Paragraph(f"Prepared by: {REPORT_AUTHOR}", styles["cover_generated"])],
        ]
        ct = Table(cover_data, colWidths=[17*cm])
        ct.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,-1),BRAND_DARK),
            ("TOPPADDING",(0,0),(-1,-1),12),("BOTTOMPADDING",(0,0),(-1,-1),12),
            ("LEFTPADDING",(0,0),(-1,-1),20),("RIGHTPADDING",(0,0),(-1,-1),20),
            ("ROUNDEDCORNERS",[8]),
        ]))
        elements.append(Spacer(1, 3*cm))
        elements.append(ct)
        elements.append(Spacer(1, 1*cm))
        elements.append(HRFlowable(width="100%", thickness=4, color=BRAND_RED))
        return elements

    def _pdf_section(self, styles, title, body):
        elements = []
        elements.append(Paragraph(title.upper(), styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        if body:
            for para in body.strip().split("\n\n"):
                elements.append(Paragraph(para.strip(), styles["body"]))
        return elements

    def _pdf_fuel_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("FUEL PERFORMANCE SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H,C,R = styles["table_header"],styles["table_cell"],styles["table_cell_right"]
        data = [
            [Paragraph("Metric",H),Paragraph("PMS (Petrol)",H),Paragraph("AGO (Diesel)",H)],
            [Paragraph("Volume Sold",C),Paragraph(_fmt(_safe(metrics,"pms_volume_total"),False)+" L",R),Paragraph(_fmt(_safe(metrics,"ago_volume_total"),False)+" L",R)],
            [Paragraph("Revenue",C),Paragraph(_fmt(_safe(metrics,"pms_revenue_total")),R),Paragraph(_fmt(_safe(metrics,"ago_revenue_total")),R)],
            [Paragraph("Combined Revenue",C),Paragraph("",C),Paragraph(_fmt(_safe(metrics,"total_fuel_revenue")),R)],
        ]
        t = Table(data, colWidths=[6*cm,5.5*cm,5.5*cm])
        t.setStyle(self._base_table_style())
        elements.append(t); elements.append(Spacer(1,0.4*cm))
        return elements

    def _pdf_payment_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("PAYMENT COLLECTION BREAKDOWN", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H,C,R = styles["table_header"],styles["table_cell"],styles["table_cell_right"]
        data = [
            [Paragraph("Payment Type",H),Paragraph("Amount",H),Paragraph("Share",H)],
            [Paragraph("Cash Collected",C),Paragraph(_fmt(_safe(metrics,"cash_collected")),R),Paragraph(_pct(_safe(metrics,"cash_percentage")),R)],
            [Paragraph("Cashless Total",C),Paragraph(_fmt(_safe(metrics,"cashless_collected")),R),Paragraph(_pct(_safe(metrics,"cashless_percentage")),R)],
            [Paragraph("  \u21b3 Plus Card",C),Paragraph(_fmt(_safe(metrics,"plus_card_total")),R),Paragraph("",C)],
            [Paragraph("  \u21b3 Visa",C),Paragraph(_fmt(_safe(metrics,"visa_total")),R),Paragraph("",C)],
            [Paragraph("  \u21b3 Credit Sales",C),Paragraph(_fmt(_safe(metrics,"credit_sales_total")),R),Paragraph("",C)],
            [Paragraph("Total Sales",C),Paragraph(_fmt(_safe(metrics,"total_sales")),R),Paragraph("100%",R)],
        ]
        t = Table(data, colWidths=[6*cm,6*cm,5*cm])
        t.setStyle(self._base_table_style())
        elements.append(t); elements.append(Spacer(1,0.4*cm))
        return elements

    def _pdf_expense_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("EXPENSE SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H,C,R = styles["table_header"],styles["table_cell"],styles["table_cell_right"]
        data = [
            [Paragraph("Item",H),Paragraph("Amount",H)],
            [Paragraph("Total Expenses",C),Paragraph(_fmt(_safe(metrics,"total_expenses")),R)],
            [Paragraph("Total Revenue",C),Paragraph(_fmt(_safe(metrics,"total_revenue")),R)],
            [Paragraph("Avg Daily Fuel Revenue",C),Paragraph(_fmt(_safe(metrics,"avg_daily_fuel_revenue")),R)],
        ]
        t = Table(data, colWidths=[9*cm,8*cm])
        t.setStyle(self._base_table_style())
        elements.append(t); elements.append(Spacer(1,0.4*cm))
        return elements

    def _pdf_reconciliation_table(self, styles, metrics):
        elements = []
        elements.append(Paragraph("CASH RECONCILIATION SUMMARY", styles["section_heading"]))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=BRAND_MID))
        elements.append(Spacer(1, 0.2*cm))
        H,C,R = styles["table_header"],styles["table_cell"],styles["table_cell_right"]
        delta_status = str(_safe(metrics,"delta_status","UNKNOWN"))
        delta_color  = SURPLUS_GREEN if delta_status=="SURPLUS" else DEFICIT_RED
        data = [
            [Paragraph("Item",H),Paragraph("Value",H)],
            [Paragraph("Cash Collected",C),Paragraph(_fmt(_safe(metrics,"cash_collected")),R)],
            [Paragraph("Total Banked",C),Paragraph(_fmt(_safe(metrics,"total_cash_banked")),R)],
            [Paragraph("Expected to Bank",C),Paragraph(_fmt(_safe(metrics,"total_cash_expected")),R)],
            [Paragraph("Net Delta",C),Paragraph(_fmt(_safe(metrics,"total_delta")),R)],
            [Paragraph("Delta Status",C),Paragraph(delta_status,R)],
            [Paragraph("Days with Anomalies",C),Paragraph(str(_safe(metrics,"anomaly_days_count")),R)],
        ]
        t = Table(data, colWidths=[9*cm,8*cm])
        style = self._base_table_style()
        t.setStyle(TableStyle([*style._cmds,("TEXTCOLOR",(1,5),(1,5),delta_color),("FONTNAME",(1,5),(1,5),"Helvetica-Bold")]))
        elements.append(t); elements.append(Spacer(1,0.4*cm))
        return elements

    def _base_table_style(self):
        return TableStyle([
            ("BACKGROUND",(0,0),(-1,0),BRAND_DARK),("TEXTCOLOR",(0,0),(-1,0),white),
            ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,0),9),
            ("TOPPADDING",(0,0),(-1,0),7),("BOTTOMPADDING",(0,0),(-1,0),7),
            ("FONTNAME",(0,1),(-1,-1),"Helvetica"),("FONTSIZE",(0,1),(-1,-1),9),
            ("TOPPADDING",(0,1),(-1,-1),5),("BOTTOMPADDING",(0,1),(-1,-1),5),
            ("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[white,BRAND_LIGHT]),
            ("GRID",(0,0),(-1,-1),0.5,BRAND_MID),("LINEABOVE",(0,0),(-1,0),1.5,BRAND_RED),
        ])

    # =========================================================
    # DOCX
    # =========================================================

    def generate_docx(self, metrics, report_sections, period_label):
        filename = _filename("StationDeck_Report", period_label, "docx")
        filepath = REPORTS_DOCX_DIR / filename
        doc = Document()
        self._docx_set_margins(doc)
        self._docx_cover(doc, period_label)
        doc.add_page_break()
        self._docx_section_heading(doc, "Executive Summary")
        self._docx_body(doc, report_sections.get("executive_summary",""))
        self._docx_section_heading(doc, "Fuel Performance Summary")
        self._docx_fuel_table(doc, metrics)
        self._docx_body(doc, report_sections.get("fuel_sales_analysis",""))
        self._docx_section_heading(doc, "Payment Collection Breakdown")
        self._docx_payment_table(doc, metrics)
        self._docx_body(doc, report_sections.get("payment_collection_analysis",""))
        self._docx_section_heading(doc, "Expense Summary")
        self._docx_expense_table(doc, metrics)
        self._docx_section_heading(doc, "Cash Reconciliation Summary")
        self._docx_reconciliation_table(doc, metrics)
        self._docx_body(doc, report_sections.get("cash_reconciliation_analysis",""))
        self._docx_footer_line(doc, period_label)
        doc.save(str(filepath))
        print(f"  \u2705 DOCX saved: {filepath}")
        return str(filepath)

    def _docx_set_margins(self, doc):
        for section in doc.sections:
            section.top_margin=Cm(2); section.bottom_margin=Cm(2)
            section.left_margin=Cm(2); section.right_margin=Cm(2)

    def _docx_cover(self, doc, period_label):
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(STATION_NAME); run.font.size=Pt(28); run.font.bold=True; run.font.color.rgb=W_DARK
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(STATION_LOCATION); run.font.size=Pt(13); run.font.color.rgb=RGBColor(0x88,0x88,0x88)
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run("OPERATIONAL PERFORMANCE REPORT"); run.font.size=Pt(14); run.font.bold=True; run.font.color.rgb=W_RED
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(period_label.upper()); run.font.size=Pt(20); run.font.bold=True; run.font.color.rgb=W_DARK
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}"); run.font.size=Pt(10); run.font.color.rgb=RGBColor(0x88,0x88,0x88)
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(f"Prepared by: {REPORT_AUTHOR}"); run.font.size=Pt(10); run.font.color.rgb=RGBColor(0x88,0x88,0x88)

    def _docx_section_heading(self, doc, title):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        run=p.add_run(title.upper()); run.font.size=Pt(12); run.font.bold=True; run.font.color.rgb=W_RED

    def _docx_body(self, doc, text):
        if not text: return
        for para in text.strip().split("\n\n"):
            p=doc.add_paragraph(); run=p.add_run(para.strip())
            run.font.size=Pt(10); run.font.color.rgb=RGBColor(0x22,0x22,0x22)
            p.paragraph_format.space_after=Pt(6)

    def _docx_fuel_table(self, doc, metrics):
        self._docx_table(doc,["Metric","PMS (Petrol)","AGO (Diesel)"],[
            ["Volume Sold",_fmt(_safe(metrics,"pms_volume_total"),False)+" L",_fmt(_safe(metrics,"ago_volume_total"),False)+" L"],
            ["Revenue",_fmt(_safe(metrics,"pms_revenue_total")),_fmt(_safe(metrics,"ago_revenue_total"))],
            ["Combined Revenue","",_fmt(_safe(metrics,"total_fuel_revenue"))],
        ])

    def _docx_payment_table(self, doc, metrics):
        self._docx_table(doc,["Payment Type","Amount","Share"],[
            ["Cash Collected",_fmt(_safe(metrics,"cash_collected")),_pct(_safe(metrics,"cash_percentage"))],
            ["Cashless Total",_fmt(_safe(metrics,"cashless_collected")),_pct(_safe(metrics,"cashless_percentage"))],
            ["  \u21b3 Plus Card",_fmt(_safe(metrics,"plus_card_total")),""],
            ["  \u21b3 Visa",_fmt(_safe(metrics,"visa_total")),""],
            ["  \u21b3 Credit Sales",_fmt(_safe(metrics,"credit_sales_total")),""],
            ["Total Sales",_fmt(_safe(metrics,"total_sales")),"100%"],
        ])

    def _docx_expense_table(self, doc, metrics):
        self._docx_table(doc,["Item","Amount"],[
            ["Total Expenses",_fmt(_safe(metrics,"total_expenses"))],
            ["Total Revenue",_fmt(_safe(metrics,"total_revenue"))],
            ["Avg Daily Fuel Revenue",_fmt(_safe(metrics,"avg_daily_fuel_revenue"))],
        ])

    def _docx_reconciliation_table(self, doc, metrics):
        self._docx_table(doc,["Item","Value"],[
            ["Cash Collected",_fmt(_safe(metrics,"cash_collected"))],
            ["Total Banked",_fmt(_safe(metrics,"total_cash_banked"))],
            ["Expected to Bank",_fmt(_safe(metrics,"total_cash_expected"))],
            ["Net Delta",_fmt(_safe(metrics,"total_delta"))],
            ["Delta Status",str(_safe(metrics,"delta_status","UNKNOWN"))],
            ["Days with Anomalies",str(_safe(metrics,"anomaly_days_count"))],
        ])

    def _docx_table(self, doc, headers, rows):
        table=doc.add_table(rows=1+len(rows),cols=len(headers)); table.style="Table Grid"
        for i,h in enumerate(table.rows[0].cells):
            h.text=""
            run=h.paragraphs[0].add_run(headers[i]); run.font.bold=True; run.font.size=Pt(9); run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            self._docx_set_cell_bg(h,"1A1A2E")
        for r_idx,row_data in enumerate(rows):
            bg="F5F5F5" if r_idx%2==0 else "FFFFFF"
            for c_idx,cell in enumerate(table.rows[r_idx+1].cells):
                cell.text=""
                run=cell.paragraphs[0].add_run(str(row_data[c_idx])); run.font.size=Pt(9)
                self._docx_set_cell_bg(cell,bg)
        doc.add_paragraph()

    def _docx_set_cell_bg(self, cell, hex_color):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
        tcPr.append(shd)

    def _docx_footer_line(self, doc, period_label):
        doc.add_paragraph(); p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(f"StationDeck Report \u00b7 {period_label} \u00b7 Generated {datetime.now().strftime('%d %B %Y, %H:%M')} \u00b7 {REPORT_AUTHOR}")
        run.font.size=Pt(8); run.font.color.rgb=RGBColor(0x88,0x88,0x88)

    # =========================================================
    # EXCEL
    # =========================================================

    def generate_xlsx(self, metrics, daily_df, period_label):
        filename = _filename("StationDeck_Report", period_label, "xlsx")
        filepath = REPORTS_XLSX_DIR / filename
        wb = openpyxl.Workbook()
        self._xlsx_summary_sheet(wb, metrics, period_label)
        self._xlsx_daily_sheet(wb, daily_df, period_label)
        if "Sheet" in wb.sheetnames:
            del wb["Sheet"]
        wb.save(str(filepath))
        print(f"  \u2705 XLSX saved: {filepath}")
        return str(filepath)

    def _xlsx_summary_sheet(self, wb, metrics, period_label):
        ws = wb.create_sheet(title="Monthly Summary")
        ws.sheet_view.showGridLines = False

        ws.merge_cells("A1:C1")
        c=ws["A1"]; c.value=f"{STATION_NAME} \u2014 {STATION_LOCATION}"
        c.font=Font(name="Calibri",bold=True,size=14,color="FFFFFF"); c.fill=_xl_fill(XL_DARK)
        c.alignment=Alignment(horizontal="center",vertical="center"); ws.row_dimensions[1].height=28

        ws.merge_cells("A2:C2")
        c=ws["A2"]; c.value=f"Operational Performance Report \u2014 {period_label}"
        c.font=Font(name="Calibri",size=11,color="FFFFFF"); c.fill=_xl_fill(XL_SUBHEAD)
        c.alignment=Alignment(horizontal="center",vertical="center"); ws.row_dimensions[2].height=22

        ws.merge_cells("A3:C3")
        c=ws["A3"]; c.value=f"Generated by {REPORT_AUTHOR} on {datetime.now().strftime('%d %B %Y, %H:%M')}"
        c.font=Font(name="Calibri",size=9,color="888888")
        c.alignment=Alignment(horizontal="center",vertical="center"); ws.row_dimensions[3].height=16

        def write_section(start_row, section_title, rows_data):
            ws.row_dimensions[start_row].height=8; start_row+=1
            ws.merge_cells(f"A{start_row}:C{start_row}")
            hdr=ws[f"A{start_row}"]; hdr.value=section_title
            hdr.font=Font(name="Calibri",bold=True,size=10,color="FFFFFF"); hdr.fill=_xl_fill(XL_RED)
            hdr.alignment=Alignment(horizontal="left",vertical="center",indent=1)
            ws.row_dimensions[start_row].height=18; start_row+=1
            _xl_write_header_row(ws,start_row,["Metric","Value",""],XL_DARK,height=16); start_row+=1
            for i,(label,value) in enumerate(rows_data):
                _xl_write_data_row(ws,start_row,[label,value,""],alt=(i%2==0),number_cols=[2])
                start_row+=1
            return start_row

        row = write_section(4,"FUEL PERFORMANCE",[
            ("PMS Volume Sold",     f"{_safe(metrics,'pms_volume_total',0):,.2f} L"),
            ("AGO Volume Sold",     f"{_safe(metrics,'ago_volume_total',0):,.2f} L"),
            ("PMS Revenue",         _fmt(_safe(metrics,"pms_revenue_total"))),
            ("AGO Revenue",         _fmt(_safe(metrics,"ago_revenue_total"))),
            ("Total Fuel Revenue",  _fmt(_safe(metrics,"total_fuel_revenue"))),
            ("Avg Daily Fuel Rev.", _fmt(_safe(metrics,"avg_daily_fuel_revenue"))),
        ])
        row = write_section(row,"NON-FUEL REVENUE",[
            ("Lubricants Revenue",  _fmt(_safe(metrics,"lubes_revenue_total"))),
            ("LPG Gas Revenue",     _fmt(_safe(metrics,"lpg_revenue_total"))),
            ("Shop Sales",          _fmt(_safe(metrics,"shop_sales_total"))),
            ("Total Non-Fuel Rev.", _fmt(_safe(metrics,"total_nonfuel_revenue"))),
            ("Total Sales",         _fmt(_safe(metrics,"total_sales"))),
            ("Total Revenue",       _fmt(_safe(metrics,"total_revenue"))),
        ])
        row = write_section(row,"PAYMENT COLLECTION",[
            ("Cash Collected",      _fmt(_safe(metrics,"cash_collected"))),
            ("Cash Share",          _pct(_safe(metrics,"cash_percentage"))),
            ("Cashless Collected",  _fmt(_safe(metrics,"cashless_collected"))),
            ("Cashless Share",      _pct(_safe(metrics,"cashless_percentage"))),
            ("Plus Card Total",     _fmt(_safe(metrics,"plus_card_total"))),
            ("Visa Total",          _fmt(_safe(metrics,"visa_total"))),
            ("Credit Sales",        _fmt(_safe(metrics,"credit_sales_total"))),
        ])
        row = write_section(row,"EXPENSES",[
            ("Total Expenses",      _fmt(_safe(metrics,"total_expenses"))),
        ])
        delta_status = str(_safe(metrics,"delta_status","UNKNOWN"))
        delta_color  = XL_GREEN if delta_status=="SURPLUS" else XL_RED
        row = write_section(row,"CASH RECONCILIATION",[
            ("Days Covered",        str(_safe(metrics,"total_days"))),
            ("Cash Collected",      _fmt(_safe(metrics,"cash_collected"))),
            ("Total Banked",        _fmt(_safe(metrics,"total_cash_banked"))),
            ("Expected to Bank",    _fmt(_safe(metrics,"total_cash_expected"))),
            ("Net Delta",           _fmt(_safe(metrics,"total_delta"))),
            ("Delta Status",        delta_status),
            ("Days with Anomalies", str(_safe(metrics,"anomaly_days_count"))),
        ])
        status_cell = ws[f"B{row-2}"]
        status_cell.font = Font(name="Calibri",bold=True,size=10,color=delta_color)
        ws.column_dimensions["A"].width=28
        ws.column_dimensions["B"].width=22
        ws.column_dimensions["C"].width=4

    def _xlsx_daily_sheet(self, wb, daily_df, period_label):
        ws = wb.create_sheet(title="Daily Records")
        ws.sheet_view.showGridLines = False
        col_count = 10
        ws.merge_cells(f"A1:{get_column_letter(col_count)}1")
        c=ws["A1"]; c.value=f"{STATION_NAME} \u2014 Daily Records \u2014 {period_label}"
        c.font=Font(name="Calibri",bold=True,size=13,color="FFFFFF"); c.fill=_xl_fill(XL_DARK)
        c.alignment=Alignment(horizontal="center",vertical="center"); ws.row_dimensions[1].height=26
        _xl_write_header_row(ws,2,["Date","PMS Vol (L)","AGO Vol (L)","PMS Rev (UGX)","AGO Rev (UGX)","Total Sales (UGX)","Cash (UGX)","Cashless (UGX)","Expenses (UGX)","Delta (UGX)"],XL_DARK,height=20)
        number_cols = list(range(2, col_count+1))
        for i,(_,row_data) in enumerate(daily_df.iterrows()):
            excel_row=i+3; alt=(i%2==0)
            def safe_val(col_name, decimals=0, _r=row_data):
                try:
                    val=_r.get(col_name,0)
                    return round(float(val),decimals) if decimals else int(float(val))
                except: return 0
            try:
                dv=row_data["date"]; date_str=dv.strftime("%d %b %Y") if hasattr(dv,"strftime") else str(dv)
            except: date_str="N/A"
            cashless_val=safe_val("cashless_total")-safe_val("total_cash")
            values=[date_str,safe_val("pms_volume",2),safe_val("ago_volume",2),safe_val("pms_revenue"),safe_val("ago_revenue"),safe_val("cashless_total"),safe_val("total_cash"),cashless_val,safe_val("total_expenses"),safe_val("delta")]
            _xl_write_data_row(ws,excel_row,values,alt=alt,number_cols=number_cols)
            dv=values[9]
            if isinstance(dv,(int,float)):
                ws[f"J{excel_row}"].font=Font(name="Calibri",size=10,color=XL_GREEN if dv>=0 else XL_RED)
        for i,w in enumerate([14,13,13,18,18,18,16,16,16,14],start=1):
            ws.column_dimensions[get_column_letter(i)].width=w
        ws.freeze_panes="A3"
'''

dest.write_text(content, encoding="utf-8")
print(f"SUCCESS: exporter.py written to {dest}")
print(f"Size: {dest.stat().st_size:,} bytes")
